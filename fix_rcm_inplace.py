"""
Modifica Domande_Aperte_RCM.docx in-place preservando le 61 immagini.
Operazioni:
  1. Blank inseriti dove mancano tra card adiacenti (raramente necessario in RCM)
  2. Per ogni immagine blank che precede il suo '---', sposta il '---' PRIMA dell'immagine
     (le immagini devono stare sul lato risposta, dopo il separatore)
  3. Blank "interni" alla domanda (tra testo Q e immagine/separatore) → rimossi
  4. Separatori duplicati → rimossi
  5. Prime righe intestazione → svuotate

Differenza da fix_flashcard_inplace.py:
  - Nessuna sezione con header maiuscolo
  - La regola "nuova domanda" richiede blank_before OBBLIGATORIO
    (evita che sotto-punti numerati nelle risposte vengano trattati come nuove card)
"""
import sys, io, re, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = r'C:\diana\POLI_new\AA_MAGISTRALE\Bionanotecnologie\Esami\Domande_Aperte_RCM.docx'
DST = r'C:\diana\POLI_new\AA_MAGISTRALE\Bionanotecnologie\Esami\Domande_Aperte_RCM_new.docx'

SEP_RE = re.compile(r'^\s*[-=—–]{1,}\s*$')
Q_RE   = re.compile(r'^\s*(\d+)\.\s+')

# Paragrafi di intestazione pura da svuotare (indici fissi delle prime 2 righe)
HEADER_PARAS = {0, 1}

def has_image(p):
    return bool(p._p.findall('.//' + qn('a:blip')))

def is_blank(p):
    return not p.text.strip() and not has_image(p)

def make_blank_para():
    return OxmlElement('w:p')

def clear_para_text(p):
    for r in p._p.findall(qn('w:r')):
        if not r.findall('.//' + qn('a:blip')):
            p._p.remove(r)

# ═══════════════════════════════════════════════════════════════════════
doc = Document(SRC)
paras = doc.paragraphs

# ── STEP 1: individua dove mancano blank tra card consecutive ────────────
expanded = []
for i, p in enumerate(paras):
    for j, sub in enumerate(p.text.split('\n')):
        expanded.append((i, j, sub.strip()))

insert_before = set()
state = 'before'
last_main_q_num = 0
blank_before = False

for (para_idx, sub_idx, t) in expanded:
    if not t:
        if sub_idx == 0:
            blank_before = True
        continue
    if SEP_RE.match(t):
        if state == 'question':
            state = 'answer'
        blank_before = False
        continue
    m = Q_RE.match(t)
    if m:
        num = int(m.group(1))
        # Per RCM: nuova domanda SOLO se blank_before (nessuna scorciatoia num==last+1)
        is_new = (num > last_main_q_num) and blank_before
        if state == 'before':
            if is_new or last_main_q_num == 0:
                state = 'question'; last_main_q_num = num
        elif is_new:
            if not blank_before:
                insert_before.add(para_idx)
            state = 'question'; last_main_q_num = num
        blank_before = False
    else:
        blank_before = False

# ── STEP 2: sposta il '---' PRIMA delle immagini blank che lo precedono ──
sep_to_move = {}   # {idx_sep: idx_img_destinazione}

for i, p in enumerate(paras):
    if has_image(p) and not p.text.strip():
        # Guarda indietro: c'è già un separatore tra Q-text e questa immagine?
        has_sep_before = False
        for k in range(i - 1, -1, -1):
            pk = paras[k]
            if is_blank(pk): continue
            if has_image(pk) and not pk.text.strip(): continue
            if SEP_RE.match(pk.text.strip()): has_sep_before = True
            break
        if not has_sep_before:
            # Guarda avanti: il prossimo non-blank non-img-only è un '---'?
            for j in range(i + 1, min(i + 6, len(paras))):
                pj = paras[j]
                if has_image(pj) and not pj.text.strip(): continue
                if is_blank(pj): continue
                if SEP_RE.match(pj.text.strip()):
                    sep_to_move[j] = i
                break

print(f'Separatori da spostare: {sorted(sep_to_move.keys())} → prima di {[sep_to_move[k] for k in sorted(sep_to_move.keys())]}')

# ── STEP 3: individua blank interni ─────────────────────────────────────
internal_blanks = set()
for i, p in enumerate(paras):
    if not is_blank(p):
        continue
    found_sep = False
    for j in range(i + 1, len(paras)):
        pj = paras[j]
        if is_blank(pj): continue
        if has_image(pj) and not pj.text.strip(): continue  # salta img-only
        if SEP_RE.match(pj.text.strip()):
            found_sep = True
        break
    if found_sep:
        internal_blanks.add(i)

# Aggiungi anche il blank immediatamente dopo ogni sep spostato
# (diventa un blank intra-risposta dopo lo spostamento del sep)
for sep_idx in sep_to_move:
    next_idx = sep_idx + 1
    if next_idx < len(paras) and is_blank(paras[next_idx]):
        internal_blanks.add(next_idx)

print(f'Blank interni da rimuovere ({len(internal_blanks)}): {sorted(internal_blanks)}')
print(f'Blank da inserire prima di: {sorted(insert_before)}')
print(f'Intestazioni da svuotare: {sorted(HEADER_PARAS)}')

# ── Applica modifiche ─────────────────────────────────────────────────────
# 1) Sposta separatori
sep_moves = []
for sep_idx, img_idx in sep_to_move.items():
    sep_moves.append((paras[sep_idx]._p, paras[img_idx]._p))

for sep_elem, img_elem in sep_moves:
    parent = sep_elem.getparent()
    if parent is None or img_elem.getparent() is None:
        continue
    parent.remove(sep_elem)
    img_elem.getparent().insert(list(img_elem.getparent()).index(img_elem), sep_elem)

# 2) Altre operazioni in ordine inverso
all_ops = []
for i in HEADER_PARAS:
    all_ops.append((i, 'clear'))
for i in insert_before:
    all_ops.append((i, 'insert_blank'))
for i in internal_blanks:
    all_ops.append((i, 'delete'))
all_ops.sort(key=lambda x: x[0], reverse=True)

for (para_idx, action) in all_ops:
    p = paras[para_idx]
    parent = p._p.getparent()
    if parent is None:
        continue
    if action == 'clear':
        clear_para_text(p)
    elif action == 'insert_blank':
        parent.insert(list(parent).index(p._p), make_blank_para())
    elif action == 'delete':
        parent.remove(p._p)

# ── STEP 4: rimuovi separatori duplicati ─────────────────────────────────
prev_was_sep = False
dups = []
for p in doc.paragraphs:
    t = p.text.strip()
    if not t and not has_image(p): prev_was_sep = False; continue
    if SEP_RE.match(t):
        if prev_was_sep: dups.append(p._p)
        prev_was_sep = True
    else:
        prev_was_sep = False

print(f'Separatori duplicati rimossi: {len(dups)}')
for elem in dups:
    elem.getparent().remove(elem)

# ── Salva ──────────────────────────────────────────────────────────────────
doc.save(DST)
print(f'\nSalvato: {DST}')

# ── Verifica image-aware ───────────────────────────────────────────────────
doc2 = Document(DST)
cur_q, cur_a, st, cards = [], [], 'q', []
for p in doc2.paragraphs:
    block = p.text.strip()
    img = bool(p._p.findall('.//' + qn('a:blip')))
    if not block and not img:
        if cur_q:
            cards.append((' '.join(cur_q)[:65], ' '.join(cur_a)[:45] if cur_a else '(vuota)'))
            cur_q, cur_a, st = [], [], 'q'
        continue
    for line in block.split('\n'):
        line = line.strip()
        if not line: continue
        if SEP_RE.match(line): st = 'a'; continue
        (cur_q if st == 'q' else cur_a).append(line)
if cur_q:
    cards.append((' '.join(cur_q)[:65], ' '.join(cur_a)[:45] if cur_a else '(vuota)'))

print(f'\nCard totali: {len(cards)}')
for i, (q, a) in enumerate(cards):
    marca = ' ⚠️ VUOTA' if a == '(vuota)' else ''
    print(f'  {i+1:2d}. {q}{marca}')
    print(f'      -> {a}')
