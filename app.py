import io
import os
import re
import json
import base64
import tempfile
import hashlib
import mimetypes
import zipfile
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional, Tuple, Any

import streamlit as st
from PIL import Image

import fitz  # PyMuPDF
from docx import Document
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches

try:
    from supabase import create_client, Client
except ImportError:
    create_client = None
    Client = Any




import random

def shuffle_quiz_if_needed(quiz):
    """Mescola il quiz se l'opzione è attiva."""
    if st.session_state.quiz_random_order:
        shuffled = quiz.copy()
        random.shuffle(shuffled)
        return shuffled
    return quiz

def shuffle_flashcards_if_needed(cards):
    """Mescola le flashcard se l'opzione è attiva."""
    if st.session_state.fc_random_order:
        shuffled = cards.copy()
        random.shuffle(shuffled)
        return shuffled
    return cards

# ============================================================
# STRUTTURE DATI
# ============================================================

@dataclass
class MediaItem:
    file_name: str
    page_or_section: str
    question_hint: str = ""


@dataclass
class QuizQuestion:
    number: str
    question: str
    options: List[str]
    correct_index: Optional[int] = None
    images: List[MediaItem] = None

    def __post_init__(self):
        if self.images is None:
            self.images = []


@dataclass
class FlashCard:
    number: str
    question: str
    definition: str
    q_images: List[MediaItem] = None
    d_images: List[MediaItem] = None

    def __post_init__(self):
        if self.q_images is None:
            self.q_images = []
        if self.d_images is None:
            self.d_images = []


# ============================================================
# UTILS
# ============================================================

LETTERS = ["A", "B", "C", "D", "E", "F"]

def normalize_text(text: str) -> str:
    text = text.replace("ﬁ", "fi").replace("ﬂ", "fl")
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def md_escape(text: str) -> str:
    """
    Mette in sicurezza i caratteri che Markdown interpreterebbe, così il testo
    del docx viene mostrato a schermo com'è (es. H_2O, a*b, C#) senza che
    underscore/asterischi attivino corsivo o grassetto indesiderati.
    """
    if not text:
        return ""
    return re.sub(r"([\\`*_{}\[\]#~|])", r"\\\1", text)


def format_paragraph_text(paragraph) -> str:
    """
    Estrae il testo da un paragrafo preservando la formattazione (grassetto,
    corsivo, elenchi) e convertendola in Markdown, dopo aver messo in sicurezza
    i caratteri speciali del testo grezzo.
    """
    if not paragraph.text.strip():
        return ""

    formatted_parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        text = md_escape(text)
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        formatted_parts.append(text)

    result = "".join(formatted_parts) if formatted_parts else md_escape(paragraph.text)

    # Marcatore per punti elenco (resi come liste Markdown)
    try:
        if paragraph.style.name.startswith("List"):
            result = f"- {result}"
    except Exception:
        pass

    return result


def option_letter_to_index(letter: str) -> Optional[int]:
    letter = letter.upper().strip()
    if letter in LETTERS:
        return LETTERS.index(letter)
    return None


def safe_name(name: str) -> str:
    name = re.sub(r"[^A-Za-z0-9_.-]+", "_", name)
    return name[:80] or "image"


# ============================================================
# LETTURA DOCX CON IMMAGINI
# ============================================================

def iter_block_items(parent):
    """
    Itera paragrafi e tabelle in ordine.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        return

    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def paragraph_images(
    paragraph: Paragraph,
    doc: Document,
    out_dir: Path,
    counter_start: int,
    source_prefix: str = "docx",
) -> Tuple[List[MediaItem], int]:
    """
    Estrae le immagini presenti in un paragrafo DOCX.

    source_prefix rende univoci i nomi delle immagini quando vengono caricati
    contemporaneamente più documenti (corrette, sbagliate e non fatte).
    """
    media = []
    counter = counter_start

    blips = paragraph._element.xpath(".//*[local-name()='blip']")
    for blip in blips:
        rel_id = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if not rel_id:
            continue

        try:
            image_part = doc.part.related_parts[rel_id]
            ext = image_part.content_type.split("/")[-1].replace("jpeg", "jpg")
            file_name = f"{safe_name(source_prefix)}_img_{counter}.{ext}"
            out_path = out_dir / file_name
            out_path.write_bytes(image_part.blob)

            media.append(
                MediaItem(
                    file_name=file_name,
                    page_or_section="DOCX",
                    question_hint=source_prefix,
                )
            )
            counter += 1
        except Exception:
            continue

    return media, counter


def extract_docx_content(uploaded_file, out_dir: Path) -> Tuple[str, List[Tuple[int, MediaItem]]]:
    """
    Restituisce testo e immagini con posizione approssimativa per riga.
    Preserva la formattazione: grassetto, corsivo, punti elenco, a capo.
    """
    doc = Document(uploaded_file)
    lines = []
    image_positions = []
    img_counter = 1

    source_name = getattr(uploaded_file, "name", "documento.docx")
    source_hash = hashlib.sha1(source_name.encode("utf-8")).hexdigest()[:8]
    source_prefix = f"{safe_name(Path(source_name).stem)}_{source_hash}"

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            imgs, img_counter = paragraph_images(
                block, doc, out_dir, img_counter, source_prefix
            )
            if text:
                lines.append(text)
            for img in imgs:
                image_positions.append((len(lines), img))
                if not text:
                    lines.append("[IMMAGINE]")
        elif isinstance(block, Table):
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))

    return normalize_text("\n".join(lines)), image_positions


# ============================================================
# LETTURA PDF CON IMMAGINI
# ============================================================

def extract_pdf_content(uploaded_file, out_dir: Path) -> Tuple[str, List[Tuple[int, MediaItem]]]:
    """
    Estrae testo e immagini da PDF.
    Per le immagini usa PyMuPDF e le assegna alla posizione testuale più vicina nella pagina.
    """
    raw = uploaded_file.read()
    pdf = fitz.open(stream=raw, filetype="pdf")

    all_lines = []
    image_positions = []
    img_counter = 1

    for page_index, page in enumerate(pdf, start=1):
        page_prefix_line = len(all_lines)
        all_lines.append(f"[PAGINA {page_index}]")

        blocks = page.get_text("dict").get("blocks", [])
        blocks_sorted = sorted(blocks, key=lambda b: (b.get("bbox", [0,0,0,0])[1], b.get("bbox", [0,0,0,0])[0]))

        for block in blocks_sorted:
            if block.get("type") == 0:
                block_lines = []
                for line in block.get("lines", []):
                    spans = [span.get("text", "") for span in line.get("spans", [])]
                    txt = "".join(spans).strip()
                    if txt:
                        block_lines.append(txt)
                if block_lines:
                    all_lines.append(" ".join(block_lines))

            elif block.get("type") == 1:
                try:
                    xref = block.get("xref")
                    if xref:
                        pix = fitz.Pixmap(pdf, xref)
                        if pix.n - pix.alpha > 3:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        file_name = f"pdf_p{page_index}_img_{img_counter}.png"
                        out_path = out_dir / file_name
                        pix.save(str(out_path))
                        pix = None
                        image_positions.append((
                            len(all_lines),
                            MediaItem(file_name=file_name, page_or_section=f"PDF pagina {page_index}")
                        ))
                        img_counter += 1
                except Exception:
                    pass

        all_lines.append("")

    return normalize_text("\n".join(all_lines)), image_positions


# ============================================================
# PARSING QUIZ
# ============================================================

QUESTION_START_RE = re.compile(
    r"""^\s*
    (?:
        (?:domanda|quesito|question|q)\s*)?
        (?P<num>\d{1,3})
        [\.\)\:\-]\s+
    """,
    re.IGNORECASE | re.VERBOSE
)

OPTION_RE = re.compile(
    r"""^\s*
    (?P<letter>[A-Fa-f])
    [\.\)\:\-]\s+
    (?P<text>.+)
    """,
    re.VERBOSE
)

ANSWER_KEY_PATTERNS = [
    re.compile(r"(?P<num>\d{1,3})\s*[\.\)\:\-]\s*(?P<letter>[A-Fa-f])\b"),
    re.compile(r"(?:domanda|quesito|question|q)\s*(?P<num>\d{1,3})\s*[:\-]?\s*(?P<letter>[A-Fa-f])\b", re.I),
]


def parse_answer_key(text: str) -> Dict[str, int]:
    """
    Cerca sezioni tipo:
    Soluzioni
    1. B
    2. C

    oppure:
    Risposte: 1B, 2C, 3A
    """
    answer_key = {}

    # Dai più peso alla parte finale del documento.
    lowered = text.lower()
    start_candidates = [
        lowered.rfind("soluzioni"),
        lowered.rfind("risposte"),
        lowered.rfind("answer key"),
        lowered.rfind("answers"),
        lowered.rfind("correzione"),
    ]
    start = max(start_candidates)

    search_area = text[start:] if start != -1 else text[-4000:]

    for pattern in ANSWER_KEY_PATTERNS:
        for m in pattern.finditer(search_area):
            num = m.group("num").strip()
            idx = option_letter_to_index(m.group("letter"))
            if idx is not None:
                answer_key[num] = idx

    # Formato compatto: 1B 2C 3A
    for m in re.finditer(r"\b(?P<num>\d{1,3})\s*(?P<letter>[A-Fa-f])\b", search_area):
        num = m.group("num")
        idx = option_letter_to_index(m.group("letter"))
        if idx is not None:
            answer_key[num] = idx

    return answer_key


def split_questions(text: str) -> List[Tuple[str, int, int]]:
    """
    Restituisce blocchi domanda come:
    (testo_blocco, riga_inizio, riga_fine)
    """
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    starts = []

    for i, line in enumerate(lines):
        if QUESTION_START_RE.match(line):
            starts.append(i)

    # Fallback: se non trova numerazione, prova a usare le righe che finiscono con ?
    if not starts:
        for i, line in enumerate(lines):
            if line.endswith("?"):
                starts.append(i)

    blocks = []
    for idx, start in enumerate(starts):
        end = starts[idx + 1] if idx + 1 < len(starts) else len(lines)
        block = "\n".join(lines[start:end])
        blocks.append((block, start, end))

    return blocks


def parse_question_block(block: str) -> Optional[QuizQuestion]:
    lines = [l.strip() for l in block.splitlines() if l.strip()]
    if not lines:
        return None

    first = lines[0]
    m = QUESTION_START_RE.match(first)
    if m:
        number = m.group("num")
        first_question_text = QUESTION_START_RE.sub("", first).strip()
    else:
        number = str(abs(hash(first)) % 100000)
        first_question_text = first

    question_lines = []
    if first_question_text:
        question_lines.append(first_question_text)

    options_dict = {}
    current_letter = None

    for line in lines[1:]:
        opt = OPTION_RE.match(line)
        if opt:
            current_letter = opt.group("letter").upper()
            options_dict[current_letter] = opt.group("text").strip()
        else:
            if current_letter and current_letter in options_dict:
                # Riga continuazione opzione
                options_dict[current_letter] += " " + line
            else:
                # Riga continuazione domanda
                question_lines.append(line)

    if not options_dict:
        return None

    option_letters = [l for l in LETTERS if l in options_dict]
    options = [options_dict[l] for l in option_letters]

    if len(options) < 2:
        return None

    return QuizQuestion(
        number=number,
        question=" ".join(question_lines).strip(),
        options=options,
        correct_index=None,
        images=[],
    )


def parse_quiz_from_text(text: str, image_positions: List[Tuple[int, MediaItem]]) -> List[QuizQuestion]:
    answer_key = parse_answer_key(text)
    blocks = split_questions(text)

    questions = []
    question_ranges = []

    for block, start, end in blocks:
        q = parse_question_block(block)
        if q:
            if q.number in answer_key and answer_key[q.number] < len(q.options):
                q.correct_index = answer_key[q.number]
            questions.append(q)
            question_ranges.append((len(questions)-1, start, end))

    # Assegna immagini alla domanda più vicina in base alla posizione.
    for line_pos, media in image_positions:
        assigned = False
        for q_idx, start, end in question_ranges:
            if start <= line_pos < end:
                questions[q_idx].images.append(media)
                assigned = True
                break

        if not assigned and question_ranges:
            # domanda precedente più vicina
            previous = None
            for q_idx, start, end in question_ranges:
                if start <= line_pos:
                    previous = q_idx
            if previous is not None:
                questions[previous].images.append(media)

    return questions




# ============================================================
# PARSER DOCX DIRETTO E ROBUSTO
# ============================================================

def flush_current_question(current: Optional[QuizQuestion], quiz: List[QuizQuestion]):
    if current is not None and len(current.options) >= 2:
        quiz.append(current)


def parse_docx_quiz_direct(uploaded_file, out_dir: Path) -> Tuple[List[QuizQuestion], str]:
    """
    Parser DOCX robusto:
    - legge paragrafi e tabelle in ordine reale;
    - collega le immagini alla domanda corrente;
    - riconosce le soluzioni finali nel formato 1. A.
    """
    doc = Document(uploaded_file)
    quiz: List[QuizQuestion] = []
    raw_lines = []
    current: Optional[QuizQuestion] = None
    current_option_index: Optional[int] = None
    img_counter = 1

    source_name = getattr(uploaded_file, "name", "documento.docx")
    source_hash = hashlib.sha1(source_name.encode("utf-8")).hexdigest()[:8]
    source_prefix = f"{safe_name(Path(source_name).stem)}_{source_hash}"

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            block_text = block.text.strip()
            imgs, img_counter = paragraph_images(
                block, doc, out_dir, img_counter, source_prefix
            )

            if block_text:
                raw_lines.append(block_text)

                q_match = QUESTION_START_RE.match(block_text)
                opt_match = OPTION_RE.match(block_text)

                if q_match:
                    flush_current_question(current, quiz)
                    number = q_match.group("num").strip()
                    question_text = QUESTION_START_RE.sub("", block_text).strip()

                    current = QuizQuestion(
                        number=number,
                        question=question_text,
                        options=[],
                        correct_index=None,
                        images=[],
                    )
                    current_option_index = None

                elif opt_match and current is not None:
                    option_text = opt_match.group("text").strip()
                    current.options.append(option_text)
                    current_option_index = len(current.options) - 1

                elif current is not None:
                    if current_option_index is not None and current.options:
                        current.options[current_option_index] += " " + block_text
                    else:
                        current.question += " " + block_text

            for img in imgs:
                raw_lines.append(f"[IMMAGINE: {img.file_name}]")
                if current is not None:
                    current.images.append(img)

        elif isinstance(block, Table):
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                row_text = " | ".join(cells)
                if row_text:
                    raw_lines.append(row_text)
                    if current is not None:
                        current.question += " " + row_text

    flush_current_question(current, quiz)

    full_text = normalize_text("\n".join(raw_lines))
    answer_key = parse_answer_key(full_text)

    for q in quiz:
        if q.number in answer_key and answer_key[q.number] < len(q.options):
            q.correct_index = answer_key[q.number]

    return quiz, full_text


# ============================================================
# IMPORT / EXPORT JSON
# ============================================================

def quiz_to_json(quiz: List[QuizQuestion]) -> str:
    return json.dumps([asdict(q) for q in quiz], ensure_ascii=False, indent=2)


def quiz_from_json(raw: str) -> List[QuizQuestion]:
    data = json.loads(raw)
    quiz = []
    for item in data:
        images = [MediaItem(**m) for m in item.get("images", [])]
        quiz.append(
            QuizQuestion(
                number=str(item.get("number", "")),
                question=item.get("question", ""),
                options=item.get("options", []),
                correct_index=item.get("correct_index", None),
                images=images,
            )
        )
    return quiz






# ============================================================
# IMPORT JSON SALVATI DALL'APP
# ============================================================

def parse_saved_state_json(raw: str) -> List[QuizQuestion]:
    """
    Legge i JSON generati dall'app per poter ricaricare:
    - domande corrette;
    - domande sbagliate;
    - domande non fatte.

    Supporta anche il formato interno prodotto da quiz_to_json.
    Nota: i JSON salvano domanda/opzioni/soluzione, ma non le immagini.
    Per mantenere le immagini bisogna ricaricare il DOCX originale.
    """
    data = json.loads(raw)
    quiz = []

    if not isinstance(data, list):
        raise ValueError("Il JSON deve contenere una lista di domande.")

    for i, item in enumerate(data, start=1):
        # Formato interno dell'app: question/options/correct_index
        if "question" in item and "options" in item:
            images = []
            for m in item.get("images", []):
                try:
                    images.append(MediaItem(**m))
                except Exception:
                    pass

            quiz.append(
                QuizQuestion(
                    number=str(item.get("number", i)),
                    question=str(item.get("question", "")),
                    options=[str(x) for x in item.get("options", [])],
                    correct_index=item.get("correct_index", None),
                    images=images,
                )
            )
            continue

        # Formato esportato dalla sezione "Salva sessione"
        if "domanda" in item and "opzioni" in item:
            options_obj = item.get("opzioni", {})
            options = []

            if isinstance(options_obj, dict):
                for letter in LETTERS:
                    if letter in options_obj:
                        options.append(str(options_obj[letter]))
            elif isinstance(options_obj, list):
                options = [str(x) for x in options_obj]

            correct_letter = str(item.get("risposta_corretta", "")).strip().upper()
            correct_index = option_letter_to_index(correct_letter)

            quiz.append(
                QuizQuestion(
                    number=str(item.get("numero_originale", item.get("numero_progressivo", i))),
                    question=str(item.get("domanda", "")),
                    options=options,
                    correct_index=correct_index,
                    images=[],
                )
            )

    return [q for q in quiz if q.question and len(q.options) >= 2]


# ============================================================
# SALVATAGGIO SESSIONE SU RICHIESTA
# ============================================================

def split_session_by_answers(quiz: List[QuizQuestion], answers: Dict[int, int]) -> Tuple[List[QuizQuestion], List[QuizQuestion], List[QuizQuestion]]:
    """
    Divide lo stato corrente in:
    - corrette;
    - sbagliate;
    - non fatte.

    Non salva nulla su disco: calcola tutto solo quando la pagina deve mostrare
    la sezione "Salva sessione".
    """
    correct_items = []
    wrong_items = []
    unanswered_items = []

    for i, q in enumerate(quiz):
        if i not in answers:
            unanswered_items.append(q)
        elif q.correct_index is not None and answers[i] == q.correct_index:
            correct_items.append(q)
        else:
            wrong_items.append(q)

    return correct_items, wrong_items, unanswered_items


def quiz_json_with_solutions(quiz_items: List[QuizQuestion]) -> bytes:
    data = []
    for i, q in enumerate(quiz_items, start=1):
        correct_letter = ""
        correct_text = ""

        if q.correct_index is not None and 0 <= q.correct_index < len(q.options):
            correct_letter = LETTERS[q.correct_index]
            correct_text = q.options[q.correct_index]

        data.append({
            "numero_progressivo": i,
            "numero_originale": q.number,
            "domanda": q.question,
            "opzioni": {
                LETTERS[j]: opt for j, opt in enumerate(q.options)
            },
            "risposta_corretta": correct_letter,
            "testo_risposta_corretta": correct_text,
        })

    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")


def quiz_txt_with_solutions(quiz_items: List[QuizQuestion]) -> bytes:
    parts = []

    for i, q in enumerate(quiz_items, start=1):
        parts.append(f"{i}. {q.question}")

        for j, opt in enumerate(q.options):
            parts.append(f"{LETTERS[j]}. {opt}")

        if q.correct_index is not None and 0 <= q.correct_index < len(q.options):
            parts.append(f"Soluzione: {LETTERS[q.correct_index]}. {q.options[q.correct_index]}")
        else:
            parts.append("Soluzione: non impostata")

        parts.append("")

    return "\n".join(parts).encode("utf-8")



def quiz_docx_with_solutions(
    quiz_items: List[QuizQuestion],
    title: str,
    media_dir: Path,
) -> bytes:
    """
    Crea un DOCX ricaricabile dall'app con:
    - domande;
    - immagini;
    - opzioni esplicite A., B., C. ...;
    - sezione Soluzioni finale.

    Le immagini vengono incluse quando sono ancora disponibili nella sessione
    corrente, cioè quando il quiz proviene dal DOCX/PDF originale.
    """
    document = Document()

    heading = document.add_heading(title, level=1)
    heading.alignment = 1

    document.add_paragraph(
        "Questo file può essere ricaricato nell'app per continuare il quiz."
    )
    document.add_paragraph("")

    solution_rows = []

    for progressive_number, q in enumerate(quiz_items, start=1):
        # Usiamo una numerazione continua, più affidabile quando il file viene ricaricato.
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{progressive_number}. {q.question}")
        run.bold = True

        # Immagini collegate alla domanda
        for media in q.images:
            image_path = media_dir / media.file_name
            if image_path.exists():
                try:
                    document.add_picture(str(image_path), width=Inches(5.8))
                except Exception:
                    # Se Word non supporta direttamente il formato, non bloccare il salvataggio.
                    document.add_paragraph(
                        f"[Immagine non inseribile automaticamente: {media.file_name}]"
                    )

        # Opzioni scritte esplicitamente, senza elenchi automatici Word
        for option_index, option in enumerate(q.options):
            document.add_paragraph(
                f"{LETTERS[option_index]}. {option}"
            )

        document.add_paragraph("")

        if q.correct_index is not None and 0 <= q.correct_index < len(q.options):
            solution_rows.append(
                (
                    progressive_number,
                    LETTERS[q.correct_index],
                    q.options[q.correct_index],
                )
            )
        else:
            solution_rows.append(
                (progressive_number, "?", "Soluzione non impostata")
            )

    document.add_page_break()
    document.add_heading("Soluzioni", level=1)

    for number, letter, correct_text in solution_rows:
        # Salva solo la lettera della soluzione.
        # Formato compatibile con il parser: 1. A
        document.add_paragraph(f"{number}. {letter}")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_download_pair(
    title: str,
    quiz_items: List[QuizQuestion],
    base_filename: str,
    icon: str,
    media_dir: Path,
):
    st.markdown(f"**{icon} {title}: {len(quiz_items)}**")

    if not quiz_items:
        st.caption("Nessuna domanda in questa sezione.")
        return

    st.download_button(
        "Scarica DOCX ricaricabile per ogni sezione",
        data=quiz_docx_with_solutions(
            quiz_items=quiz_items,
            title=title,
            media_dir=media_dir,
        ),
        file_name=f"{base_filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        width='stretch',
    )



def render_save_session_panel(quiz: List[QuizQuestion], answers: Dict[int, int], media_dir: Path, previous_correct: List[QuizQuestion], previous_wrong: List[QuizQuestion], previous_unanswered: List[QuizQuestion]):
    correct_items, wrong_items, unanswered_items = combine_previous_and_current_state(
        quiz,
        answers,
        previous_correct,
        previous_wrong,
        previous_unanswered,
    )

    st.header("💾 Salva sessione")
    st.write(
        "I file scaricati includono anche lo storico delle sessioni precedenti. "
        "Contengono domande, opzioni, soluzioni e le immagini disponibili; "
        "puoi ricaricarli direttamente nell’app per continuare."
    )

    metric_cols = st.columns(3)
    with metric_cols[0]:
        st.metric("✅ Corrette", len(correct_items))
    with metric_cols[1]:
        st.metric("❌ Sbagliate", len(wrong_items))
    with metric_cols[2]:
        st.metric("⏳ Non fatte", len(unanswered_items))

    st.divider()

    save_cols = st.columns(3)

    with save_cols[0]:
        render_download_pair(
            "Domande corrette",
            correct_items,
            "quiz_domande_corrette_con_soluzioni",
            "✅",
            media_dir,
        )

    with save_cols[1]:
        render_download_pair(
            "Domande sbagliate",
            wrong_items,
            "quiz_domande_sbagliate_con_soluzioni",
            "❌",
            media_dir,
        )

    with save_cols[2]:
        render_download_pair(
            "Domande non fatte",
            unanswered_items,
            "quiz_domande_non_fatte_con_soluzioni",
            "⏳",
            media_dir,
        )

    if wrong_items:
        st.info("Puoi ricaricare il DOCX delle domande sbagliate per rifare solo quelle.")
    elif len(answers) == len(quiz) and quiz:
        st.success("Ottimo: non ci sono domande sbagliate nella sessione corrente.")




# ============================================================
# MERGE SESSIONI PRECEDENTI
# ============================================================

def stable_question_key(q: QuizQuestion) -> str:
    qtext = re.sub(r"\s+", " ", q.question.strip().lower())
    opts = [re.sub(r"\s+", " ", x.strip().lower()) for x in q.options]
    return qtext + "||" + "||".join(opts)


def merge_unique_questions(*groups: List[QuizQuestion]) -> List[QuizQuestion]:
    result = {}
    order = []
    for group in groups:
        for q in group:
            key = stable_question_key(q)
            if key not in result:
                result[key] = q
                order.append(key)
            else:
                old = result[key]
                if not old.images and q.images:
                    result[key] = q
                elif old.correct_index is None and q.correct_index is not None:
                    result[key] = q
    return [result[k] for k in order]


def remove_questions(base: List[QuizQuestion], removed: List[QuizQuestion]) -> List[QuizQuestion]:
    keys = {stable_question_key(q) for q in removed}
    return [q for q in base if stable_question_key(q) not in keys]


def classify_uploaded_session_docx(file_name: str) -> Optional[str]:
    name = file_name.lower()
    if "corrett" in name:
        return "correct"
    if "sbagliat" in name or "error" in name:
        return "wrong"
    if "non_fatt" in name or "nonfatt" in name:
        return "unanswered"
    return None


def combine_previous_and_current_state(
    current_quiz: List[QuizQuestion],
    current_answers: Dict[int, int],
    previous_correct: List[QuizQuestion],
    previous_wrong: List[QuizQuestion],
    previous_unanswered: List[QuizQuestion],
) -> Tuple[List[QuizQuestion], List[QuizQuestion], List[QuizQuestion]]:
    current_correct, current_wrong, current_unanswered = split_session_by_answers(
        current_quiz, current_answers
    )

    correct = merge_unique_questions(previous_correct, current_correct)
    wrong = merge_unique_questions(previous_wrong, current_wrong)
    unanswered = merge_unique_questions(previous_unanswered, current_unanswered)

    wrong = remove_questions(wrong, current_correct)
    unanswered = remove_questions(unanswered, current_correct)

    correct = remove_questions(correct, current_wrong)
    unanswered = remove_questions(unanswered, current_wrong)

    wrong = remove_questions(wrong, correct)
    unanswered = remove_questions(unanswered, correct)

    return correct, wrong, unanswered

def select_quiz_by_mode(
    mode: str,
    original_quiz: List[QuizQuestion],
    previous_correct: List[QuizQuestion],
    previous_wrong: List[QuizQuestion],
    previous_unanswered: List[QuizQuestion],
) -> List[QuizQuestion]:

    if mode == "Non fatte":
        return merge_unique_questions(previous_unanswered)

    if mode == "Sbagliate":
        return merge_unique_questions(previous_wrong)

    if mode == "Corrette":
        return merge_unique_questions(previous_correct)

    if mode == "Sbagliate + non fatte":
        return merge_unique_questions(
            previous_wrong,
            previous_unanswered,
        )

    if mode == "Tutte":
        return merge_unique_questions(
            previous_correct,
            previous_wrong,
            previous_unanswered,
            original_quiz,
        )

    return merge_unique_questions(previous_unanswered)



# ============================================================
# CLOUD E GESTIONE SESSIONI
# ============================================================

class NamedBytesIO(io.BytesIO):
    def __init__(self, data: bytes, name: str):
        super().__init__(data)
        self.name = name


def get_users_config() -> Dict[str, Dict[str, str]]:
    """
    Legge l'elenco utenti da secrets.toml, sezione [users.<nome>].
    Ogni utente ha una password e una propria cartella "root" nel bucket,
    così i dati di persone diverse non si mescolano.
    """
    try:
        raw_users = st.secrets["users"]
    except Exception:
        return {}

    users = {}
    for username, cfg in raw_users.items():
        users[username] = {
            "password": str(cfg.get("password", "")),
            "root": safe_name(str(cfg.get("root", username))),
        }
    return users


def render_login_gate(users: Dict[str, Dict[str, str]]) -> bool:
    """
    Mostra un form di login se ci sono utenti configurati.
    Ritorna True se l'utente è autenticato (o se non è configurato nessun
    utente, per restare compatibile con l'uso a utente singolo).
    """
    if not users:
        return True

    if st.session_state.get("auth_user"):
        return True

    st.title("☁️ Quiz Cloud Interattivi")
    st.subheader("🔐 Accedi")
    with st.form("login_form"):
        username = st.text_input("Nome utente")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Accedi", type="primary")

    if submitted:
        user = users.get(username)
        if user and password and user["password"] == password:
            st.session_state.auth_user = username
            st.session_state.auth_root = user["root"]
            st.rerun()
        else:
            st.error("Nome utente o password non corretti.")

    return False


def get_cloud_settings() -> Optional[Dict[str, str]]:
    try:
        cfg = st.secrets["cloud"]
        return {
            "url": str(cfg["supabase_url"]),
            "key": str(cfg["supabase_key"]),
            "bucket": str(cfg.get("bucket", "quiz-cloud")),
            "root": safe_name(str(cfg.get("root", "utente"))),
        }
    except Exception:
        return None


@st.cache_resource
def get_supabase_client(url: str, key: str):
    if create_client is None:
        raise RuntimeError(
            "Pacchetto 'supabase' non installato. Esegui: pip install -r requirements.txt"
        )
    return create_client(url, key)


def get_cloud_context():
    settings = get_cloud_settings()
    if not settings:
        return None, None

    client = get_supabase_client(settings["url"], settings["key"])
    return client, settings


def ensure_bucket(client, bucket: str):
    try:
        client.storage.get_bucket(bucket)
    except Exception:
        client.storage.create_bucket(
            bucket,
            options={
                "public": False,
                "allowed_mime_types": [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "application/pdf",
                    "application/json",
                    "application/zip",
                    "text/plain",
                ],
            },
        )


def cloud_path(*parts: str) -> str:
    # Alcune chiamate passano un "part" che è già un percorso composto
    # (es. base = cloud_path(root, subject, folder)). safe_name() sostituisce
    # qualsiasi carattere non alfanumerico con "_", quindi se applicata
    # all'intero percorso distrugge anche gli "/" interni, appiattendo le
    # sottocartelle alla radice del bucket. Per questo dividiamo prima ogni
    # part sugli "/" e sanifichiamo ogni segmento singolarmente.
    segments = []
    for part in parts:
        part = str(part).strip().strip("/")
        if not part:
            continue
        for segment in part.split("/"):
            segment = segment.strip()
            if segment:
                segments.append(safe_name(segment))
    return "/".join(segments)


def cloud_list(client, bucket: str, path: str = "") -> List[Dict[str, Any]]:
    result = client.storage.from_(bucket).list(
        path,
        {
            "limit": 1000,
            "offset": 0,
            "sortBy": {"column": "name", "order": "asc"},
        },
    )
    return result or []


def cloud_upload_bytes(
    client,
    bucket: str,
    path: str,
    data: bytes,
    mime_type: str,
):
    return client.storage.from_(bucket).upload(
        path=path,
        file=data,
        file_options={
            "content-type": mime_type,
            "cache-control": "3600",
            "upsert": "true",
        },
    )


def cloud_download_bytes(client, bucket: str, path: str) -> bytes:
    data = client.storage.from_(bucket).download(path)
    if isinstance(data, bytes):
        return data
    if hasattr(data, "content"):
        return bytes(data.content)
    return bytes(data)


def create_subject_folder(client, settings: Dict[str, str], subject: str):
    subject = safe_name(subject)
    if not subject:
        raise ValueError("Inserisci un nome valido per la materia.")

    path = cloud_path(settings["root"], subject, ".keep")
    cloud_upload_bytes(client, settings["bucket"], path, b"", "text/plain")


def list_subjects(client, settings: Dict[str, str]) -> List[str]:
    rows = cloud_list(client, settings["bucket"], settings["root"])
    subjects = []

    for row in rows:
        name = str(row.get("name", "")).strip()
        if not name or name.startswith("."):
            continue
        # Le cartelle virtuali Supabase sono restituite senza id/metadata file.
        if row.get("id") is None or row.get("metadata") is None:
            subjects.append(name)

    return sorted(set(subjects), key=str.lower)


def list_subject_contents(
    client,
    settings: Dict[str, str],
    subject: str,
) -> Tuple[List[str], List[str]]:
    prefix = cloud_path(settings["root"], subject)
    rows = cloud_list(client, settings["bucket"], prefix)

    files = []
    folders = []

    for row in rows:
        name = str(row.get("name", "")).strip()
        if not name or name.startswith("."):
            continue

        if row.get("id") is None or row.get("metadata") is None:
            folders.append(name)
        elif name.lower().endswith((".docx", ".pdf")):
            files.append(name)

    return sorted(files, key=str.lower), sorted(folders, key=str.lower)


def upload_original_quiz(
    client,
    settings: Dict[str, str],
    subject: str,
    uploaded_file,
):
    file_name = safe_name(uploaded_file.name)
    if not file_name.lower().endswith((".docx", ".pdf")):
        raise ValueError("Sono ammessi soltanto file DOCX e PDF.")

    data = uploaded_file.getvalue()
    mime = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if file_name.lower().endswith(".docx")
        else "application/pdf"
    )
    path = cloud_path(settings["root"], subject, file_name)
    cloud_upload_bytes(client, settings["bucket"], path, data, mime)
    return file_name

def delete_cloud_file(
    client,
    settings,
    subject,
    file_name,
):
    path = cloud_path(
        settings["root"],
        subject,
        file_name,
    )

    client.storage.from_(
        settings["bucket"]
    ).remove([path])

def delete_session_folder(
    client,
    settings,
    subject,
    folder_name,
):
    base = cloud_path(
        settings["root"],
        subject,
        folder_name,
    )

    files = [
        cloud_path(base, "corrette.docx"),
        cloud_path(base, "sbagliate.docx"),
        cloud_path(base, "non_fatte.docx"),
    ]

    client.storage.from_(
        settings["bucket"]
    ).remove(files)

def parse_file_bytes(
    data: bytes,
    file_name: str,
    media_dir: Path,
) -> Tuple[List[QuizQuestion], str]:
    file_obj = NamedBytesIO(data, file_name)

    if file_name.lower().endswith(".docx"):
        return parse_docx_quiz_direct(file_obj, media_dir)

    if file_name.lower().endswith(".pdf"):
        text, image_positions = extract_pdf_content(file_obj, media_dir)
        return parse_quiz_from_text(text, image_positions), text

    raise ValueError("Formato file non supportato.")


def load_original_from_cloud(
    client,
    settings: Dict[str, str],
    subject: str,
    file_name: str,
    media_dir: Path,
) -> Tuple[List[QuizQuestion], str]:
    path = cloud_path(settings["root"], subject, file_name)
    data = cloud_download_bytes(client, settings["bucket"], path)
    return parse_file_bytes(data, file_name, media_dir)


SESSION_FILES = {
    "correct": "corrette.docx",
    "wrong": "sbagliate.docx",
    "unanswered": "non_fatte.docx",
}


def load_session_folder(
    client,
    settings: Dict[str, str],
    subject: str,
    folder_name: str,
    media_dir: Path,
) -> Tuple[List[QuizQuestion], List[QuizQuestion], List[QuizQuestion]]:
    base = cloud_path(settings["root"], subject, folder_name)
    loaded: Dict[str, List[QuizQuestion]] = {
        "correct": [],
        "wrong": [],
        "unanswered": [],
    }

    for key, file_name in SESSION_FILES.items():
        path = cloud_path(base, file_name)
        try:
            data = cloud_download_bytes(client, settings["bucket"], path)
            quiz, _ = parse_file_bytes(data, file_name, media_dir)
            loaded[key] = quiz
        except Exception:
            loaded[key] = []

    return loaded["correct"], loaded["wrong"], loaded["unanswered"]

def get_quiz_status(client, settings, subject, file_name):
    try:
        session_folder = (
            Path(file_name).stem + "_sessione"
        )

        stats = load_stats_json(
            client,
            settings,
            subject,
            session_folder,
        )

        corrette = stats["corrette"]
        sbagliate = stats["sbagliate"]
        non_fatte = stats["non_fatte"]

        if corrette == 0 and sbagliate == 0:
            return "⚪"

        if sbagliate > 0:
            return "🔴"

        if non_fatte > 0:
            return "🟡"

        return "🟢"

    except:
        return "⚪"

def make_session_folder_name(quiz_name: str) -> str:
    stem = Path(quiz_name).stem
    return f"{safe_name(stem)}_sessione"


def save_session_to_cloud(
    client,
    settings: Dict[str, str],
    subject: str,
    session_folder: str,
    correct_items: List[QuizQuestion],
    wrong_items: List[QuizQuestion],
    unanswered_items: List[QuizQuestion],
    media_dir: Path,
):
    base = cloud_path(settings["root"], subject, session_folder)
    docs = {
        "corrette.docx": quiz_docx_with_solutions(
            correct_items, "Domande corrette", media_dir
        ),
        "sbagliate.docx": quiz_docx_with_solutions(
            wrong_items, "Domande sbagliate", media_dir
        ),
        "non_fatte.docx": quiz_docx_with_solutions(
            unanswered_items, "Domande non fatte", media_dir
        ),
    }

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    for file_name, data in docs.items():
        cloud_upload_bytes(
            client,
            settings["bucket"],
            cloud_path(base, file_name),
            data,
            mime,
        )

    metadata = {
        "subject": subject,
        "session_folder": session_folder,
        "correct_count": len(correct_items),
        "wrong_count": len(wrong_items),
        "unanswered_count": len(unanswered_items),
    }
    cloud_upload_bytes(
        client,
        settings["bucket"],
        cloud_path(base, "sessione.json"),
        json.dumps(metadata, ensure_ascii=False, indent=2).encode("utf-8"),
        "application/json",
    )

    # stats.json: usato dall'indicatore di stato (pallino) accanto al quiz.
    # Le chiavi devono corrispondere a quelle lette da load_stats_json /
    # get_quiz_status (corrette / sbagliate / non_fatte).
    stats = {
        "corrette": len(correct_items),
        "sbagliate": len(wrong_items),
        "non_fatte": len(unanswered_items),
    }
    cloud_upload_bytes(
        client,
        settings["bucket"],
        cloud_path(base, "stats.json"),
        json.dumps(stats, ensure_ascii=False, indent=2).encode("utf-8"),
        "application/json",
    )


def session_backup_zip(
    correct_items: List[QuizQuestion],
    wrong_items: List[QuizQuestion],
    unanswered_items: List[QuizQuestion],
    media_dir: Path,
) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "corrette.docx",
            quiz_docx_with_solutions(correct_items, "Domande corrette", media_dir),
        )
        zf.writestr(
            "sbagliate.docx",
            quiz_docx_with_solutions(wrong_items, "Domande sbagliate", media_dir),
        )
        zf.writestr(
            "non_fatte.docx",
            quiz_docx_with_solutions(
                unanswered_items, "Domande non fatte", media_dir
            ),
        )
    return output.getvalue()


def clear_active_quiz_state():
    st.session_state.quiz = []
    st.session_state.original_quiz = []
    st.session_state.previous_correct = []
    st.session_state.previous_wrong = []
    st.session_state.previous_unanswered = []
    st.session_state.answers = {}
    st.session_state.editing_done = False
    st.session_state.active_subject = ""
    st.session_state.active_quiz_name = ""
    st.session_state.active_session_folder = ""
    st.session_state.loaded_from_session = False


def start_original_quiz(
    quiz: List[QuizQuestion],
    subject: str,
    quiz_name: str,
):
    st.session_state.original_quiz = quiz.copy()
    st.session_state.quiz = quiz.copy()
    st.session_state.previous_correct = []
    st.session_state.previous_wrong = []
    st.session_state.previous_unanswered = quiz.copy()
    st.session_state.answers = {}
    st.session_state.editing_done = False
    st.session_state.active_subject = subject
    st.session_state.active_quiz_name = quiz_name
    st.session_state.active_session_folder = make_session_folder_name(quiz_name)
    st.session_state.loaded_from_session = False
    st.session_state.quiz_mode = "Tutte"


def start_saved_session(
    correct: List[QuizQuestion],
    wrong: List[QuizQuestion],
    unanswered: List[QuizQuestion],
    subject: str,
    folder_name: str,
):
    all_questions = merge_unique_questions(correct, wrong, unanswered)
    st.session_state.previous_correct = correct
    st.session_state.previous_wrong = wrong
    st.session_state.previous_unanswered = unanswered
    st.session_state.original_quiz = all_questions
    st.session_state.quiz = all_questions
    st.session_state.answers = {}
    st.session_state.editing_done = False
    st.session_state.active_subject = subject
    st.session_state.active_quiz_name = folder_name.removesuffix("_sessione")
    st.session_state.active_session_folder = folder_name
    st.session_state.loaded_from_session = True
    st.session_state.quiz_mode = "Non fatte" if unanswered else "Sbagliate"


def current_combined_state():
    return combine_previous_and_current_state(
        st.session_state.quiz,
        st.session_state.answers,
        st.session_state.previous_correct,
        st.session_state.previous_wrong,
        st.session_state.previous_unanswered,
    )

# ============================================================
# FLASHCARD - PARSING
# ============================================================

FLASHCARD_SEP_RE = re.compile(
    r"^\s*[-=]{3,}\s*$"  # linee divisorie tipo --- o ===
)

def parse_flashcards_from_docx(uploaded_file, out_dir: Path) -> Tuple[List[FlashCard], str]:
    """
    Legge un DOCX con flashcard nel formato:
        1. Domanda (testo o immagine)
        ---
        Definizione (testo o immagine)
        ---
        2. Prossima domanda ...

    Oppure il formato semplice (alternanza paragrafo domanda / paragrafo definizione).
    Le immagini vengono collegate alla domanda o alla definizione in base alla posizione.
    """
    doc = Document(uploaded_file)
    cards: List[FlashCard] = []
    raw_lines = []

    source_name = getattr(uploaded_file, "name", "flashcard.docx")
    source_hash = hashlib.sha1(source_name.encode("utf-8")).hexdigest()[:8]
    source_prefix = f"fc_{safe_name(Path(source_name).stem)}_{source_hash}"

    img_counter = 1

    # Stati del parser: "question" o "definition"
    current_number = 0
    current_question = []
    current_q_imgs: List[MediaItem] = []
    current_definition = []
    current_d_imgs: List[MediaItem] = []
    state = "question"  # stiamo leggendo la domanda o la definizione?

    def flush_card():
        nonlocal current_number
        q_text = "\n\n".join(s for s in current_question if s).strip()
        d_text = "\n\n".join(s for s in current_definition if s).strip()
        if q_text or current_q_imgs:
            current_number += 1
            cards.append(FlashCard(
                number=str(current_number),
                question=q_text,
                definition=d_text,
                q_images=list(current_q_imgs),
                d_images=list(current_d_imgs),
            ))

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            block_text = block.text.strip()
            block_fmt = format_paragraph_text(block)
            imgs, img_counter = paragraph_images(
                block, doc, out_dir, img_counter, source_prefix
            )

            if block_text:
                raw_lines.append(block_text)

            # Le immagini si attaccano allo stato attivo all'inizio del
            # paragrafo, indipendentemente da separatori/domande individuati
            # più sotto nello stesso paragrafo (vedi nota sotto).
            for img in imgs:
                if state == "question":
                    current_q_imgs.append(img)
                else:
                    current_d_imgs.append(img)

            # Un paragrafo Word può contenere più "righe logiche" separate da
            # un a capo "soft" (Shift+Invio) invece che da un nuovo paragrafo:
            # es. domanda, "---" e risposta scritte tutte di seguito nello
            # stesso paragrafo. Le elaboriamo riga per riga così
            # domanda/separatore/risposta vengono riconosciuti comunque,
            # anche quando non sono andati a capo "per davvero".
            text_lines = block_text.split("\n") if block_text else []
            fmt_lines = block_fmt.split("\n") if block_fmt else []
            while len(fmt_lines) < len(text_lines):
                fmt_lines.append("")
            while len(text_lines) < len(fmt_lines):
                text_lines.append("")

            for line_text, line_fmt in zip(text_lines, fmt_lines):
                line_text = line_text.strip()
                line_fmt = line_fmt.strip()
                if not line_text and not line_fmt:
                    continue

                # Controlla se è un separatore tra domanda e definizione
                if FLASHCARD_SEP_RE.match(line_text) or line_text.lower() in ("---", "===", "***"):
                    if state == "question":
                        state = "definition"
                    else:
                        # Separatore tra una card e la successiva
                        flush_card()
                        current_question.clear()
                        current_q_imgs.clear()
                        current_definition.clear()
                        current_d_imgs.clear()
                        state = "question"
                    continue

                # Controlla se è l'inizio di una nuova domanda numerata
                q_match = QUESTION_START_RE.match(line_text)
                if q_match:
                    flush_card()
                    current_question.clear()
                    current_q_imgs.clear()
                    current_definition.clear()
                    current_d_imgs.clear()
                    state = "question"
                    q_text = QUESTION_START_RE.sub("", line_fmt).strip()
                    if q_text:
                        current_question.append(q_text)
                    continue

                # Accumula testo nella sezione corrente
                if state == "question":
                    if line_fmt:
                        current_question.append(line_fmt)
                else:
                    if line_fmt:
                        current_definition.append(line_fmt)

        elif isinstance(block, Table):
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                row_text = " | ".join(cells)
                if row_text:
                    raw_lines.append(row_text)
                    row_fmt = " | ".join(md_escape(c) for c in cells)
                    if state == "question":
                        current_question.append(row_fmt)
                    else:
                        current_definition.append(row_fmt)

    flush_card()

    return cards, normalize_text("\n".join(raw_lines))


def parse_flashcards_from_pdf(uploaded_file, out_dir: Path) -> Tuple[List[FlashCard], str]:
    """
    Legge un PDF cercando il formato:
        1. Domanda
        ---
        Definizione

    Oppure riconosce domanda numerata + paragrafo successivo come definizione.
    """
    text, image_positions = extract_pdf_content(uploaded_file, out_dir)
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    cards: List[FlashCard] = []
    img_by_line = {}
    for line_pos, media in image_positions:
        img_by_line.setdefault(line_pos, []).append(media)

    i = 0
    card_num = 0
    while i < len(lines):
        line = lines[i]
        q_match = QUESTION_START_RE.match(line)
        if q_match:
            card_num += 1
            q_text = QUESTION_START_RE.sub("", line).strip()
            q_imgs = list(img_by_line.get(i, []))
            i += 1

            # Accumula righe non-numerata fino al separatore o prossima numerata
            extra_q = []
            while i < len(lines) and not QUESTION_START_RE.match(lines[i]) and not FLASHCARD_SEP_RE.match(lines[i]):
                if lines[i].lower() in ("---", "===", "***"):
                    break
                extra_q.append(lines[i])
                q_imgs += img_by_line.get(i, [])
                i += 1

            if extra_q:
                q_text = (q_text + " " + " ".join(extra_q)).strip()

            # Salta separatore
            if i < len(lines) and (FLASHCARD_SEP_RE.match(lines[i]) or lines[i].lower() in ("---", "===", "***")):
                i += 1

            # Accumula definizione
            d_lines = []
            d_imgs = []
            while i < len(lines) and not QUESTION_START_RE.match(lines[i]) and not FLASHCARD_SEP_RE.match(lines[i]):
                if lines[i].lower() in ("---", "===", "***"):
                    i += 1
                    break
                d_lines.append(lines[i])
                d_imgs += img_by_line.get(i, [])
                i += 1

            cards.append(FlashCard(
                number=str(card_num),
                question=q_text,
                definition=" ".join(d_lines).strip(),
                q_images=q_imgs,
                d_images=d_imgs,
            ))
        else:
            i += 1

    return cards, text


def parse_flashcard_bytes(data: bytes, file_name: str, media_dir: Path) -> Tuple[List[FlashCard], str]:
    file_obj = NamedBytesIO(data, file_name)
    if file_name.lower().endswith(".docx"):
        return parse_flashcards_from_docx(file_obj, media_dir)
    if file_name.lower().endswith(".pdf"):
        return parse_flashcards_from_pdf(file_obj, media_dir)
    raise ValueError("Formato file non supportato.")


def flashcards_to_json(cards: List[FlashCard]) -> bytes:
    from dataclasses import asdict
    return json.dumps([asdict(c) for c in cards], ensure_ascii=False, indent=2).encode("utf-8")


def flashcards_from_json(raw: bytes) -> List[FlashCard]:
    data = json.loads(raw)
    cards = []
    for item in data:
        q_images = [MediaItem(**m) for m in item.get("q_images", [])]
        d_images = [MediaItem(**m) for m in item.get("d_images", [])]
        cards.append(FlashCard(
            number=str(item.get("number", "")),
            question=item.get("question", ""),
            definition=item.get("definition", ""),
            q_images=q_images,
            d_images=d_images,
        ))
    return cards


# ============================================================
# FLASHCARD - CLOUD
# ============================================================

def upload_flashcard_file(client, settings, subject, uploaded_file):
    file_name = safe_name(uploaded_file.name)
    if not file_name.lower().endswith((".docx", ".pdf")):
        raise ValueError("Sono ammessi soltanto file DOCX e PDF.")
    data = uploaded_file.getvalue()
    mime = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if file_name.lower().endswith(".docx") else "application/pdf"
    )
    path = cloud_path(settings["root"], subject, "flashcard", file_name)
    cloud_upload_bytes(client, settings["bucket"], path, data, mime)
    return file_name


def list_flashcard_files(client, settings, subject):
    prefix = cloud_path(settings["root"], subject, "flashcard")
    rows = cloud_list(client, settings["bucket"], prefix)
    files = []
    for row in rows:
        name = str(row.get("name", "")).strip()
        if name and not name.startswith(".") and name.lower().endswith((".docx", ".pdf")):
            files.append(name)
    return sorted(files, key=str.lower)


def delete_flashcard_file(client, settings, subject, file_name):
    path = cloud_path(settings["root"], subject, "flashcard", file_name)
    client.storage.from_(settings["bucket"]).remove([path])


def load_flashcard_session(client, settings, subject, file_stem) -> dict:
    """Carica lo stato sessione flashcard (note e conosco/da_studiare)."""
    path = cloud_path(settings["root"], subject, "flashcard", f"{file_stem}_stato.json")
    try:
        data = cloud_download_bytes(client, settings["bucket"], path)
        return json.loads(data)
    except Exception:
        return {"conosco": [], "da_studiare": []}


def save_flashcard_session(client, settings, subject, file_stem, stato: dict):
    path = cloud_path(settings["root"], subject, "flashcard", f"{file_stem}_stato.json")
    data = json.dumps(stato, ensure_ascii=False, indent=2).encode("utf-8")
    cloud_upload_bytes(client, settings["bucket"], path, data, "application/json")


def load_flashcard_from_cloud(client, settings, subject, file_name, media_dir):
    path = cloud_path(settings["root"], subject, "flashcard", file_name)
    data = cloud_download_bytes(client, settings["bucket"], path)
    return parse_flashcard_bytes(data, file_name, media_dir)


def load_stats_json(
    client,
    settings,
    subject,
    session_folder,
):
    try:
        stats_path = cloud_path(
            settings["root"],
            subject,
            session_folder,
            "stats.json",
        )

        response = (
            client.storage
            .from_(settings["bucket"])
            .download(stats_path)
        )

        data = json.loads(response.decode("utf-8"))

        return {
            "corrette": int(data.get("corrette", 0)),
            "sbagliate": int(data.get("sbagliate", 0)),
            "non_fatte": int(data.get("non_fatte", 0)),
        }

    except Exception:
        return {
            "corrette": 0,
            "sbagliate": 0,
            "non_fatte": 0,
        }

def save_stats_json(
    client,
    settings,
    subject,
    session_folder,
    stats,
):
    stats_path = cloud_path(
        settings["root"],
        subject,
        session_folder,
        "stats.json",
    )

    payload = json.dumps(
        stats,
        ensure_ascii=False,
        indent=2,
    ).encode("utf-8")

    client.storage.from_(
        settings["bucket"]
    ).upload(
        stats_path,
        payload,
        {"upsert": "true"},
    )

# ============================================================
# STREAMLIT APP
# ============================================================

st.set_page_config(
    page_title="Quiz Cloud Interattivi",
    page_icon="☁️",
    layout="wide",
)

st.markdown(
    """
<style>
.block-container {max-width: 1280px; padding-top: 1.4rem;}
[data-testid="stMetric"] {
    background: rgba(127,127,127,0.08);
    border: 1px solid rgba(127,127,127,0.18);
    padding: 12px;
    border-radius: 14px;
}
.quiz-card {
    border: 1px solid rgba(127,127,127,0.22);
    border-radius: 16px;
    padding: 16px;
    margin-bottom: 10px;
}
.cloud-path {
    font-size: 0.9rem;
    opacity: 0.75;
}
</style>
""",
    unsafe_allow_html=True,
)

# Session state
defaults = {
    "work_dir": tempfile.mkdtemp(prefix="quiz_cloud_"),
    "quiz": [],
    "original_quiz": [],
    "previous_correct": [],
    "previous_wrong": [],
    "previous_unanswered": [],
    "answers": {},
    "editing_done": False,
    "quiz_mode": "Non fatte",
    "active_subject": "",
    "active_quiz_name": "",
    "active_session_folder": "",
    "loaded_from_session": False,
    "selected_subject": "",
    "raw_text": "",
    # Flashcard state
    "fc_cards": [],
    "fc_file_name": "",
    "fc_subject": "",
    "fc_mode": "studia",         # "studia" o "prova_tu"
    "fc_study_filter": "Tutte",  # Tutte / Da studiare / Conosco
    "fc_conosco": [],            # numeri carte marcate "conosco"
    "fc_da_studiare": [],        # numeri carte marcate "da studiare"
    "fc_flipped": {},            # {card_index: bool} carta girata o no
    "fc_prove_risposte": {},     # {card_index: str} risposte scritte
    "quiz_random_order": False,  # Ordine casuale per quiz
    "fc_random_order": False,    # Ordine casuale per flashcard
}
for key, value in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value

work_dir = Path(st.session_state.work_dir)
media_dir = work_dir / "media"
media_dir.mkdir(parents=True, exist_ok=True)

configured_users = get_users_config()
if not render_login_gate(configured_users):
    st.stop()

client, cloud_settings = get_cloud_context()
cloud_ready = client is not None and cloud_settings is not None

if cloud_ready and st.session_state.get("auth_root"):
    # Ogni utente ha la sua cartella: i dati non si mescolano tra persone diverse.
    cloud_settings = dict(cloud_settings)
    cloud_settings["root"] = st.session_state.auth_root

st.title("☁️ Quiz Cloud Interattivi")
st.caption(
    "Materie, quiz originali e sessioni salvate in un unico spazio. "
    "Le sessioni vengono aggiornate con un solo pulsante."
)

with st.sidebar:
    if st.session_state.get("auth_user"):
        st.caption(f"👤 Accesso come **{st.session_state.auth_user}**")
        if st.button("🚪 Esci", width='stretch'):
            st.session_state.pop("auth_user", None)
            st.session_state.pop("auth_root", None)
            st.rerun()
        st.divider()

    st.header("Stato cloud")
    if cloud_ready:
        try:
            ensure_bucket(client, cloud_settings["bucket"])
            st.success("Cloud collegato")
            st.caption(f"Bucket: {cloud_settings['bucket']}")
            st.caption(f"Spazio: {cloud_settings['root']}")
        except Exception as exc:
            cloud_ready = False
            st.error(f"Cloud non disponibile: {exc}")
    else:
        st.warning("Cloud non configurato")
        st.caption(
            "Compila `.streamlit/secrets.toml` seguendo il file "
            "`secrets.example.toml`."
        )

    st.divider()
    
    if st.session_state.active_subject or st.session_state.fc_cards:
        st.header("📚 Elemento attivo")
        if st.session_state.active_subject:
            st.write(f"**Materia:** {st.session_state.active_subject}")
            st.write(f"**Quiz:** {st.session_state.active_quiz_name}")
            if st.session_state.active_session_folder:
                st.write(f"**Cartella:** {st.session_state.active_session_folder}")
        
        if st.session_state.fc_cards:
            st.write(f"**Materia FC:** {st.session_state.fc_subject}")
            st.write(f"**Mazzo:** {st.session_state.fc_file_name}")

        save_disabled = not cloud_ready or not (
            (st.session_state.active_subject and st.session_state.active_session_folder)
            or (st.session_state.fc_cards and st.session_state.fc_file_name)
        )
        if st.button("💾 Salva sessione", type="primary", width='stretch', disabled=save_disabled):
            try:
                if st.session_state.active_subject and st.session_state.active_session_folder:
                    combined_correct, combined_wrong, combined_unanswered = current_combined_state()
                    save_session_to_cloud(
                        client,
                        cloud_settings,
                        st.session_state.active_subject,
                        st.session_state.active_session_folder,
                        combined_correct,
                        combined_wrong,
                        combined_unanswered,
                        media_dir,
                    )
                    st.session_state.previous_correct = combined_correct
                    st.session_state.previous_wrong = combined_wrong
                    st.session_state.previous_unanswered = combined_unanswered
                if st.session_state.fc_cards and st.session_state.fc_file_name:
                    fc_file_stem = safe_name(Path(st.session_state.fc_file_name).stem)
                    save_flashcard_session(
                        client,
                        cloud_settings,
                        st.session_state.fc_subject,
                        fc_file_stem,
                        {
                            "conosco": st.session_state.fc_conosco,
                            "da_studiare": st.session_state.fc_da_studiare,
                        },
                    )
                st.success("Sessione salvata nel cloud.")
            except Exception as exc:
                st.error(f"Salvataggio non riuscito: {exc}")
        if save_disabled and cloud_ready:
            st.caption("Apri un quiz o un mazzo flashcard per poter salvare.")

        if st.button("🏠 Torna alla Home", type="secondary", width='stretch'):
            clear_active_quiz_state()
            st.session_state.requested_tab = "cloud"
            st.session_state.fc_cards = []
            st.session_state.fc_file_name = ""
            st.session_state.fc_subject = ""
            st.session_state.fc_conosco = []
            st.session_state.fc_da_studiare = []
            st.session_state.fc_flipped = {}
            st.session_state.fc_prove_risposte = {}
            st.rerun()
        
        st.divider()
        if st.button("❌ Azzera elemento attivo", width='stretch', key="clear_active"):
            clear_active_quiz_state()
            st.rerun()
    else:
        st.info("Carica un quiz o una flashcard per iniziare.")


TAB_OPTIONS = {
    "cloud": "☁️ Le mie materie",
    "quiz": "📝 Quiz interattivo",
    "locale": "📤 Importazione locale",
    "flashcard": "🃏 Flashcard",
}
_tab_keys = list(TAB_OPTIONS.keys())

# Il widget con key="active_tab" non può essere modificato dopo essere stato
# istanziato in questo stesso run. Per cambiare sezione da un bottone, il
# bottone scrive in "requested_tab"; qui, PRIMA di creare il widget,
# applichiamo la richiesta ad active_tab e la consumiamo.
if st.session_state.get("requested_tab") in _tab_keys:
    st.session_state.active_tab = st.session_state.pop("requested_tab")
st.session_state.pop("requested_tab", None)

if st.session_state.get("active_tab") not in _tab_keys:
    st.session_state.active_tab = "cloud"

selected_tab = st.radio(
    "Sezione",
    options=_tab_keys,
    format_func=lambda k: TAB_OPTIONS[k],
    key="active_tab",
    horizontal=True,
    label_visibility="collapsed",
)

# ------------------------------------------------------------------
# CLOUD BROWSER
# ------------------------------------------------------------------
if selected_tab == "cloud":
    if not cloud_ready:
        st.info(
            "Configura Supabase per usare cartelle e salvataggi persistenti. "
            "Nel frattempo puoi provare l'app dalla scheda Importazione locale."
        )
    else:
        ensure_bucket(client, cloud_settings["bucket"])

        st.subheader("Materie")
        create_col, refresh_col = st.columns([4, 1])
        with create_col:
            new_subject = st.text_input(
                "Nuova materia",
                placeholder="es. Bionanotecnologie",
                label_visibility="collapsed",
            )
        with refresh_col:
            create_clicked = st.button("＋ Crea materia", type="primary", width='stretch')

        if create_clicked:
            try:
                create_subject_folder(client, cloud_settings, new_subject)
                st.success(f"Materia '{safe_name(new_subject)}' creata.")
                st.rerun()
            except Exception as exc:
                st.error(f"Impossibile creare la materia: {exc}")

        subjects = list_subjects(client, cloud_settings)

        if not subjects:
            st.info("Non ci sono ancora materie. Creane una qui sopra.")
        else:
            selected_subject = st.selectbox(
                "Apri materia",
                subjects,
                index=(
                    subjects.index(st.session_state.selected_subject)
                    if st.session_state.selected_subject in subjects
                    else 0
                ),
            )
            st.session_state.selected_subject = selected_subject

            st.markdown(
                f'<div class="cloud-path">☁️ / {selected_subject}</div>',
                unsafe_allow_html=True,
            )

            st.divider()

            # ---- Upload affiancato: quiz | flashcard ----
            col_quiz_up, col_fc_up = st.columns(2)

            with col_quiz_up:
                st.subheader("📝 Carica quiz")
                with st.expander("📋 Formato consigliato per i quiz"):
                    st.code(
"""1. Testo della domanda
A. prima opzione
B. seconda opzione
C. terza opzione
D. quarta opzione

Soluzioni

1. B
2. C""", language="text")
                    st.markdown("""
**Immagini** — inseriscile nel DOCX dopo il testo della domanda. Il numero con il punto è obbligatorio; le opzioni iniziano con `A.` `B.` `C.` `D.`; la sezione `Soluzioni` va alla fine.
""")
                quiz_upload = st.file_uploader(
                    "DOCX o PDF (anche multipli)",
                    type=["docx", "pdf"],
                    accept_multiple_files=True,
                    key=f"cloud_upload_{selected_subject}",
                )
                if st.button(
                    "Carica quiz nel cloud",
                    type="primary",
                    disabled=not quiz_upload,
                    key=f"btn_upload_quiz_{selected_subject}",
                    width='stretch',
                ):
                    try:
                        names = []
                        for file in quiz_upload:
                            names.append(upload_original_quiz(client, cloud_settings, selected_subject, file))
                        st.success(f"Caricati {len(names)} quiz.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"Caricamento non riuscito: {exc}")

            with col_fc_up:
                st.subheader("🃏 Carica flashcard")
                with st.expander("📋 Formato consigliato per le flashcard"):
                    st.code(
"""1. Termine o domanda
[immagine opzionale nel DOCX]
---
Definizione o risposta
[immagine opzionale nel DOCX]

2. Altro termine
---
Altra definizione""", language="text")
                    st.markdown("""
**Immagini** — inseriscile direttamente nel DOCX:
- Prima del `---` → appare con la **domanda**
- Dopo il `---` → appare con la **definizione**

Il numero con il punto è obbligatorio. Il `---` su una riga da sola separa domanda e definizione.
""")
                fc_upload_cloud = st.file_uploader(
                    "DOCX o PDF",
                    type=["docx", "pdf"],
                    accept_multiple_files=False,
                    key=f"fc_cloud_upload_{selected_subject}",
                )
                if st.button(
                    "Carica flashcard nel cloud",
                    type="primary",
                    disabled=not fc_upload_cloud,
                    key=f"btn_upload_fc_{selected_subject}",
                    width='stretch',
                ):
                    try:
                        upload_flashcard_file(client, cloud_settings, selected_subject, fc_upload_cloud)
                        st.success(f"'{fc_upload_cloud.name}' caricato.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"Caricamento non riuscito: {exc}")

            st.divider()

            # ---- Contenuto della materia ----
            files, folders = list_subject_contents(client, cloud_settings, selected_subject)
            fc_files = list_flashcard_files(client, cloud_settings, selected_subject)

            st.subheader("Contenuto della materia")

            if not files and not folders and not fc_files:
                st.info("La materia è vuota.")
            else:
                # Quiz originali
                if files:
                    st.markdown("**📝 Quiz**")
                    for file_name in files:
                        has_session = make_session_folder_name(file_name) in folders
                        c1, c2, c3, c4 = st.columns([5, 1, 1, 1])
                        with c1:
                            status = get_quiz_status(client, cloud_settings, selected_subject, file_name)
                            st.markdown(f"{status} 📄 **{file_name}**")
                            st.caption(
                                "Quiz con sessione salvata: 'Avvia' riprende da dove eri rimasto."
                                if has_session
                                else "Quiz originale"
                            )
                        with c2:
                            if st.button("Avvia", key=f"open_file_{selected_subject}_{file_name}", width='stretch'):
                                try:
                                    for old in media_dir.glob("*"):
                                        old.unlink(missing_ok=True)
                                    quiz, raw_text = load_original_from_cloud(client, cloud_settings, selected_subject, file_name, media_dir)
                                    session_name = make_session_folder_name(file_name)
                                    if session_name in folders:
                                        # Esiste già una sessione salvata per questo quiz:
                                        # la ricarichiamo per non perdere corrette/sbagliate/non fatte.
                                        correct, wrong, unanswered = load_session_folder(
                                            client, cloud_settings, selected_subject, session_name, media_dir
                                        )
                                        start_saved_session(correct, wrong, unanswered, selected_subject, session_name)
                                        st.session_state.original_quiz = quiz
                                        st.session_state.quiz_mode = "Tutte"
                                        st.session_state.quiz = select_quiz_by_mode(
                                            "Tutte",
                                            st.session_state.original_quiz,
                                            st.session_state.previous_correct,
                                            st.session_state.previous_wrong,
                                            st.session_state.previous_unanswered,
                                        )
                                    else:
                                        start_original_quiz(quiz, selected_subject, file_name)
                                    st.session_state.raw_text = raw_text
                                    st.session_state.requested_tab = "quiz"  # Vai direttamente al quiz
                                    st.rerun()
                                except Exception as exc:
                                    st.error(f"Errore: {exc}")
                        with c3:
                            if st.button("🗑️", key=f"delete_file_{selected_subject}_{file_name}", width='stretch'):
                                try:
                                    delete_cloud_file(client, cloud_settings, selected_subject, file_name)
                                    st.success("Eliminato.")
                                    st.rerun()
                                except Exception as exc:
                                    st.error(exc)
                        with c4:
                            if has_session:
                                if st.button("♻️", key=f"reset_session_{selected_subject}_{file_name}", width='stretch', help="Azzera il progresso salvato per questo quiz"):
                                    try:
                                        session_name = make_session_folder_name(file_name)
                                        delete_session_folder(client, cloud_settings, selected_subject, session_name)
                                        st.success("Progresso azzerato.")
                                        st.rerun()
                                    except Exception as exc:
                                        st.error(exc)

                # Flashcard
                if fc_files:
                    st.markdown("**🃏 Flashcard**")
                    for cf_name in fc_files:
                        cf1, cf2, cf3 = st.columns([6, 1, 1])
                        with cf1:
                            icon = "🃏" if cf_name.endswith(".json") else "📄"
                            st.markdown(f"{icon} **{cf_name}**")
                            st.caption("Mazzo flashcard")
                        with cf2:
                            if st.button("Studia", key=f"fc_open_{selected_subject}_{cf_name}", width='stretch'):
                                try:
                                    for old in media_dir.glob("fc_*"):
                                        old.unlink(missing_ok=True)
                                    if cf_name.endswith(".json"):
                                        path = cloud_path(cloud_settings["root"], selected_subject, "flashcard", cf_name)
                                        raw = cloud_download_bytes(client, cloud_settings["bucket"], path)
                                        cards = flashcards_from_json(raw)
                                    else:
                                        cards, _ = load_flashcard_from_cloud(client, cloud_settings, selected_subject, cf_name, media_dir)
                                    st.session_state.fc_cards = cards
                                    st.session_state.fc_file_name = cf_name
                                    st.session_state.fc_subject = selected_subject
                                    st.session_state.fc_flipped = {}
                                    st.session_state.fc_prove_risposte = {}
                                    file_stem = safe_name(Path(cf_name).stem)
                                    stato = load_flashcard_session(client, cloud_settings, selected_subject, file_stem)
                                    st.session_state.fc_conosco = stato.get("conosco", [])
                                    st.session_state.fc_da_studiare = stato.get("da_studiare", [])
                                    st.session_state.requested_tab = "flashcard"
                                    st.success(f"Caricate {len(cards)} flashcard.")
                                    st.rerun()
                                except Exception as exc:
                                    st.error(f"Errore: {exc}")
                        with cf3:
                            if st.button("🗑️", key=f"fc_del_{selected_subject}_{cf_name}", width='stretch'):
                                try:
                                    if cf_name.endswith(".json"):
                                        path = cloud_path(cloud_settings["root"], selected_subject, "flashcard", cf_name)
                                        client.storage.from_(cloud_settings["bucket"]).remove([path])
                                    else:
                                        delete_flashcard_file(client, cloud_settings, selected_subject, cf_name)
                                    st.success("Eliminato.")
                                    st.rerun()
                                except Exception as exc:
                                    st.error(exc)
# ------------------------------------------------------------------
# LOCAL IMPORT, kept as fallback and to preserve previous workflow
# ------------------------------------------------------------------
if selected_tab == "locale":
    st.subheader("Apri un quiz dal dispositivo")
    st.write(
        "Questa modalità non salva automaticamente nel cloud. "
        "È utile per verificare un file prima di caricarlo."
    )

    local_file = st.file_uploader(
        "Scegli DOCX o PDF",
        type=["docx", "pdf"],
        key="local_quiz_file",
    )
    local_subject = st.text_input(
        "Materia temporanea",
        value="Locale",
        key="local_subject",
    )

    if st.button(
        "Apri quiz locale",
        type="primary",
        disabled=local_file is None,
    ):
        try:
            for old in media_dir.glob("*"):
                old.unlink(missing_ok=True)
            quiz, raw_text = parse_file_bytes(
                local_file.getvalue(),
                local_file.name,
                media_dir,
            )
            start_original_quiz(quiz, local_subject, local_file.name)
            st.session_state.raw_text = raw_text
            st.session_state.requested_tab = "quiz"
            st.success(f"Caricate {len(quiz)} domande.")
            st.rerun()
        except Exception as exc:
            st.error(f"Errore nell'importazione: {exc}")

# ------------------------------------------------------------------
# QUIZ
# ------------------------------------------------------------------
if selected_tab == "quiz":
    if not st.session_state.original_quiz:
        st.info(
            "Apri un file o una cartella sessione dalla scheda "
            "“Le mie materie”."
        )
    else:
        st.subheader(
            f"{st.session_state.active_subject} · "
            f"{st.session_state.active_quiz_name}"
        )

        historical_cols = st.columns(3)
        with historical_cols[0]:
            st.metric(
                "✅ Corrette salvate",
                len(st.session_state.previous_correct),
            )
        with historical_cols[1]:
            st.metric(
                "❌ Sbagliate salvate",
                len(st.session_state.previous_wrong),
            )
        with historical_cols[2]:
            st.metric(
                "⏳ Non fatte salvate",
                len(st.session_state.previous_unanswered),
            )

        mode_options = [
            "Non fatte",
            "Sbagliate",
            "Corrette",
            "Sbagliate + non fatte",
            "Tutte",
        ]
        default_index = (
            mode_options.index(st.session_state.quiz_mode)
            if st.session_state.quiz_mode in mode_options
            else 0
        )

        st.session_state.quiz_mode = st.radio(
            "Quali domande vuoi svolgere?",
            mode_options,
            index=default_index,
            horizontal=True,
        )

        launch_col, review_col = st.columns([2, 1])
        with launch_col:
            launch_quiz = st.button(
                "▶ Avvia modalità scelta",
                type="primary",
                width='stretch',
            )
        with review_col:
            review_first = st.checkbox(
                "Controlla domande prima di iniziare",
                value=False,
            )

        if launch_quiz:
            selected = select_quiz_by_mode(
                st.session_state.quiz_mode,
                st.session_state.original_quiz,
                st.session_state.previous_correct,
                st.session_state.previous_wrong,
                st.session_state.previous_unanswered,
            )
            if not selected:
                st.warning(
                    f"Non ci sono domande nella modalità "
                    f"“{st.session_state.quiz_mode}”."
                )
            else:
                st.session_state.quiz = selected
                st.session_state.answers = {}
                st.session_state.editing_done = not review_first
                st.rerun()

        quiz: List[QuizQuestion] = st.session_state.quiz

        if quiz and not st.session_state.editing_done:
            st.divider()
            st.header("Revisione del quiz")
            st.caption(
                "Modifica solo ciò che è stato riconosciuto male. "
                "Le immagini restano collegate alle domande."
            )

            for i, q in enumerate(quiz):
                with st.expander(
                    f"Domanda {i + 1} · originale {q.number}",
                    expanded=i == 0,
                ):
                    q.question = st.text_area(
                        "Testo domanda",
                        q.question,
                        key=f"edit_question_{i}",
                    )

                    updated_options = []
                    for j, option in enumerate(q.options):
                        updated_options.append(
                            st.text_input(
                                f"Opzione {LETTERS[j]}",
                                option,
                                key=f"edit_option_{i}_{j}",
                            )
                        )
                    q.options = updated_options

                    current_idx = (
                        q.correct_index
                        if q.correct_index is not None
                        and q.correct_index < len(q.options)
                        else 0
                    )
                    q.correct_index = st.selectbox(
                        "Soluzione",
                        range(len(q.options)),
                        index=current_idx,
                        format_func=lambda idx, q=q: (
                            f"{LETTERS[idx]}. {q.options[idx][:100]}"
                        ),
                        key=f"edit_correct_{i}",
                    )

                    for media in q.images:
                        img_path = media_dir / media.file_name
                        if img_path.exists():
                            st.image(
                                str(img_path),
                                caption=media.page_or_section,
                                width='stretch',
                            )

            if st.button(
                "Conferma e inizia",
                type="primary",
                width='stretch',
            ):
                st.session_state.quiz = quiz
                st.session_state.answers = {}
                st.session_state.editing_done = True
                st.rerun()

        elif quiz and st.session_state.editing_done:
            current_correct, current_wrong, current_unanswered = (
                split_session_by_answers(
                    quiz, st.session_state.answers
                )
            )

            p1, p2, p3, p4 = st.columns(4)
            with p1:
                st.metric("Totali", len(quiz))
            with p2:
                st.metric("✅ Corrette ora", len(current_correct))
            with p3:
                st.metric("❌ Sbagliate ora", len(current_wrong))
            with p4:
                st.metric("⏳ Non fatte ora", len(current_unanswered))

            if quiz:
                st.progress(len(st.session_state.answers) / len(quiz))

            st.divider()

            for i, q in enumerate(quiz):
                st.markdown(
                    f'<div class="quiz-card"><strong>'
                    f'Domanda {i + 1}</strong></div>',
                    unsafe_allow_html=True,
                )
                st.write(q.question)

                for media in q.images:
                    img_path = media_dir / media.file_name
                    if img_path.exists():
                        st.image(
                            str(img_path),
                            caption=media.page_or_section,
                            width='stretch',
                        )

                answer = st.radio(
                    "Scegli la risposta",
                    range(len(q.options)),
                    format_func=lambda idx, q=q: (
                        f"{LETTERS[idx]}. {q.options[idx]}"
                    ),
                    key=f"answer_{i}",
                    index=st.session_state.answers.get(i),
                )

                if answer is not None:
                    st.session_state.answers[i] = answer
                    if q.correct_index is None:
                        st.warning("Soluzione non impostata.")
                    elif answer == q.correct_index:
                        st.success("✅ Risposta corretta!")
                    else:
                        st.error("❌ Risposta sbagliata.")
                        st.write(
                            f"**Soluzione:** "
                            f"{LETTERS[q.correct_index]}. "
                            f"{q.options[q.correct_index]}"
                        )

                st.divider()

            combined_correct, combined_wrong, combined_unanswered = (
                current_combined_state()
            )

            st.header("💾 Salva la sessione")
            st.write(
                "Un solo salvataggio aggiorna nel cloud i tre documenti: "
                "corrette, sbagliate e non fatte."
            )

            s1, s2, s3 = st.columns(3)
            with s1:
                st.metric("✅ Corrette totali", len(combined_correct))
            with s2:
                st.metric("❌ Sbagliate totali", len(combined_wrong))
            with s3:
                st.metric("⏳ Non fatte totali", len(combined_unanswered))

            save_cloud_disabled = (
                not cloud_ready
                or not st.session_state.active_subject
                or not st.session_state.active_session_folder
            )

            save_col, backup_col = st.columns(2)

            with save_col:
                if st.button(
                    "☁️ Salva/aggiorna tutto nel cloud",
                    type="primary",
                    width='stretch',
                    disabled=save_cloud_disabled,
                ):
                    try:
                        save_session_to_cloud(
                            client,
                            cloud_settings,
                            st.session_state.active_subject,
                            st.session_state.active_session_folder,
                            combined_correct,
                            combined_wrong,
                            combined_unanswered,
                            media_dir,
                        )
                        st.session_state.previous_correct = combined_correct
                        st.session_state.previous_wrong = combined_wrong
                        st.session_state.previous_unanswered = (
                            combined_unanswered
                        )
                        st.success(
                            "Sessione aggiornata nella stessa cartella cloud."
                        )
                    except Exception as exc:
                        st.error(f"Salvataggio non riuscito: {exc}")

            with backup_col:
                backup = session_backup_zip(
                    combined_correct,
                    combined_wrong,
                    combined_unanswered,
                    media_dir,
                )
                st.download_button(
                    "⬇️ Scarica backup completo",
                    data=backup,
                    file_name=(
                        f"{safe_name(st.session_state.active_quiz_name)}"
                        "_sessione.zip"
                    ),
                    mime="application/zip",
                    width='stretch',
                )

            if not cloud_ready:
                st.caption(
                    "Configura il cloud per salvare senza dipendere dal PC. "
                    "Il backup ZIP resta comunque disponibile."
                )

# ------------------------------------------------------------------
# FLASHCARD TAB
# ------------------------------------------------------------------
def fc_clear_active():
    st.session_state.fc_cards = []
    st.session_state.fc_file_name = ""
    st.session_state.fc_subject = ""
    st.session_state.fc_conosco = []
    st.session_state.fc_da_studiare = []
    st.session_state.fc_flipped = {}
    st.session_state.fc_prove_risposte = {}


def fc_render_study_session():
    all_cards: List[FlashCard] = st.session_state.fc_cards
    conosco_set = set(st.session_state.fc_conosco)
    da_studiare_set = set(st.session_state.fc_da_studiare)

    st.subheader(
        f"{st.session_state.fc_subject} · {st.session_state.fc_file_name} · {len(all_cards)} carte"
    )
    m1, m2, m3 = st.columns(3)
    with m1:
        st.metric("✅ Conosco", len(conosco_set))
    with m2:
        st.metric("📖 Da studiare", len(da_studiare_set))
    with m3:
        st.metric("⚪ Non ancora viste", len(all_cards) - len(conosco_set) - len(da_studiare_set))

    mode_col, filter_col, random_col = st.columns(3)
    with mode_col:
        fc_mode = st.radio("Modalità", ["🔍 Studia", "✏️ Prova tu"], horizontal=True, key="fc_mode_radio")
    with filter_col:
        fc_filter_val = st.radio("Mostra", ["Tutte", "Non viste", "Da studiare", "Conosco"], horizontal=True, key="fc_filter_radio")
    with random_col:
        st.session_state.fc_random_order = st.checkbox("🔀 Casuale", value=False, help="Mostra le carte in ordine casuale.")

    def fc_filter(card: FlashCard) -> bool:
        if fc_filter_val == "Tutte": return True
        if fc_filter_val == "Conosco": return card.number in conosco_set
        if fc_filter_val == "Da studiare": return card.number in da_studiare_set
        if fc_filter_val == "Non viste": return card.number not in conosco_set and card.number not in da_studiare_set
        return True

    visible_cards = [c for c in all_cards if fc_filter(c)]

    if not visible_cards:
        st.info("Nessuna carta corrisponde al filtro selezionato.")
    else:
        st.caption(f"Mostrando {len(visible_cards)} su {len(all_cards)} carte")
        is_study = fc_mode.startswith("🔍")

        for idx, card in enumerate(visible_cards):
            card_key = f"fc_{card.number}_{idx}"
            is_flipped = st.session_state.fc_flipped.get(card_key, False)

            if card.number in conosco_set:
                border_color = "#2ecc71"; stato_emoji = "✅"
            elif card.number in da_studiare_set:
                border_color = "#e74c3c"; stato_emoji = "📖"
            else:
                border_color = "rgba(127,127,127,0.22)"; stato_emoji = "⚪"

            st.markdown(
                f'<div style="border:2px solid {border_color};border-radius:16px;padding:16px;margin-bottom:16px;">',
                unsafe_allow_html=True,
            )
            st.markdown(f"**{stato_emoji} Carta #{card.number}**")
            if card.question:
                st.write(card.question)
            for img in card.q_images:
                ip = media_dir / img.file_name
                if ip.exists():
                    st.image(str(ip), width='stretch')

            if is_study:
                flip_label = "🔽 Mostra definizione" if not is_flipped else "🔼 Nascondi definizione"
                if st.button(flip_label, key=f"flip_{card_key}"):
                    st.session_state.fc_flipped[card_key] = not is_flipped
                    st.rerun()
                if is_flipped:
                    st.markdown("---")
                    if card.definition:
                        st.info(card.definition)
                    for img in card.d_images:
                        ip = media_dir / img.file_name
                        if ip.exists():
                            st.image(str(ip), width='stretch')
                    bc1, bc2 = st.columns(2)
                    with bc1:
                        if st.button("✅ Conosco", key=f"fc_know_{card_key}", width='stretch'):
                            st.session_state.fc_conosco = list((set(st.session_state.fc_conosco) | {card.number}) - da_studiare_set)
                            st.session_state.fc_da_studiare = [n for n in st.session_state.fc_da_studiare if n != card.number]
                            st.session_state.fc_flipped[card_key] = False
                            st.rerun()
                    with bc2:
                        if st.button("📖 Da studiare", key=f"fc_study_{card_key}", width='stretch'):
                            st.session_state.fc_da_studiare = list((set(st.session_state.fc_da_studiare) | {card.number}) - conosco_set)
                            st.session_state.fc_conosco = [n for n in st.session_state.fc_conosco if n != card.number]
                            st.session_state.fc_flipped[card_key] = False
                            st.rerun()
            else:
                st.markdown("---")
                user_answer = st.text_area("La tua risposta", value=st.session_state.fc_prove_risposte.get(card_key, ""),
                    key=f"fc_answer_{card_key}", height=80, placeholder="Scrivi qui la tua risposta...")
                st.session_state.fc_prove_risposte[card_key] = user_answer
                if st.button("👁 Vedi definizione", key=f"fc_reveal_{card_key}"):
                    st.session_state.fc_flipped[card_key] = True
                    st.rerun()
                if is_flipped:
                    st.markdown("**Definizione:**")
                    if card.definition:
                        st.success(card.definition)
                    for img in card.d_images:
                        ip = media_dir / img.file_name
                        if ip.exists():
                            st.image(str(ip), width='stretch')
                    st.markdown("*Valuta la tua risposta:*")
                    pt1, pt2, pt3 = st.columns(3)
                    with pt1:
                        if st.button("✅ Sapevo", key=f"fc_pt_know_{card_key}", width='stretch'):
                            st.session_state.fc_conosco = list((set(st.session_state.fc_conosco) | {card.number}) - da_studiare_set)
                            st.session_state.fc_da_studiare = [n for n in st.session_state.fc_da_studiare if n != card.number]
                            st.session_state.fc_flipped[card_key] = False
                            st.rerun()
                    with pt2:
                        if st.button("📖 Da rivedere", key=f"fc_pt_study_{card_key}", width='stretch'):
                            st.session_state.fc_da_studiare = list((set(st.session_state.fc_da_studiare) | {card.number}) - conosco_set)
                            st.session_state.fc_conosco = [n for n in st.session_state.fc_conosco if n != card.number]
                            st.session_state.fc_flipped[card_key] = False
                            st.rerun()
                    with pt3:
                        if st.button("🔄 Riprova", key=f"fc_pt_retry_{card_key}", width='stretch'):
                            st.session_state.fc_prove_risposte[card_key] = ""
                            st.session_state.fc_flipped[card_key] = False
                            st.rerun()

            st.markdown("</div>", unsafe_allow_html=True)

    st.divider()
    sv_col, az_col = st.columns(2)
    with sv_col:
        save_disabled = not cloud_ready or not st.session_state.fc_file_name
        if st.button("☁️ Salva sessione flashcard", type="primary", width='stretch', disabled=save_disabled):
            try:
                file_stem = safe_name(Path(st.session_state.fc_file_name).stem)
                save_flashcard_session(client, cloud_settings, st.session_state.fc_subject, file_stem,
                    {"conosco": st.session_state.fc_conosco, "da_studiare": st.session_state.fc_da_studiare})
                st.success("✅ Sessione salvata nel cloud. Torno alla home...")
                # Torna alla home dopo 1 secondo
                import time
                time.sleep(1)
                fc_clear_active()
                st.rerun()
            except Exception as exc:
                st.error(f"Errore: {exc}")
        if save_disabled and not cloud_ready:
            st.caption("Cloud non configurato — progressi non salvati automaticamente.")
    with az_col:
        if st.button("🔄 Azzera progressi", width='stretch'):
            st.session_state.fc_conosco = []
            st.session_state.fc_da_studiare = []
            st.session_state.fc_flipped = {}
            st.session_state.fc_prove_risposte = {}
            st.rerun()
    if st.button("↩️ Torna alla libreria flashcard", width='stretch'):
        fc_clear_active()
        st.rerun()


if selected_tab == "flashcard":
    st.subheader("🃏 Flashcard")

    if st.session_state.fc_cards:
        fc_render_study_session()
    else:
        fc_sub_crea, fc_sub_cloud, fc_sub_file = st.tabs(
            ["✏️ Crea nuove", "☁️ Libreria cloud", "📤 Carica da file"]
        )

        with fc_sub_crea:
            st.markdown("Crea le tue flashcard direttamente qui. Ogni carta ha una **domanda** e una **definizione**.")
            fc_new_subject = st.text_input("Materia / raccolta", value=st.session_state.selected_subject or "Generale", key="fc_new_subject")
            fc_new_name = st.text_input("Nome mazzo", value="Mio mazzo", key="fc_new_name")
            st.divider()

            if "fc_editor_n" not in st.session_state:
                st.session_state.fc_editor_n = 3
            col_add, col_rem = st.columns(2)
            with col_add:
                if st.button("➕ Aggiungi carta", width='stretch'):
                    st.session_state.fc_editor_n += 1
                    st.rerun()
            with col_rem:
                if st.session_state.fc_editor_n > 1:
                    if st.button("➖ Rimuovi ultima", width='stretch'):
                        st.session_state.fc_editor_n -= 1
                        st.rerun()
            st.caption(f"{st.session_state.fc_editor_n} carte")

            for i in range(st.session_state.fc_editor_n):
                with st.expander(f"Carta #{i+1}", expanded=(i < 3)):
                    st.text_input("Domanda / termine", key=f"fc_ed_q_{i}", placeholder="Es. Cos'è la mitosi?")
                    st.text_area("Definizione / risposta", key=f"fc_ed_d_{i}", height=100, placeholder="Es. Processo di divisione cellulare...")

            st.divider()
            btn_s, btn_c = st.columns(2)
            with btn_s:
                if st.button("▶ Inizia a studiare", type="primary", width='stretch'):
                    cards = [FlashCard(number=str(i+1),
                                      question=st.session_state.get(f"fc_ed_q_{i}", "").strip(),
                                      definition=st.session_state.get(f"fc_ed_d_{i}", "").strip())
                             for i in range(st.session_state.fc_editor_n)
                             if st.session_state.get(f"fc_ed_q_{i}", "").strip() or st.session_state.get(f"fc_ed_d_{i}", "").strip()]
                    if not cards:
                        st.warning("Aggiungi almeno una carta.")
                    else:
                        st.session_state.fc_cards = cards
                        st.session_state.fc_file_name = safe_name(fc_new_name) + ".flashcard"
                        st.session_state.fc_subject = fc_new_subject
                        st.session_state.fc_conosco = []
                        st.session_state.fc_da_studiare = []
                        st.session_state.fc_flipped = {}
                        st.session_state.fc_prove_risposte = {}
                        st.rerun()
            with btn_c:
                if st.button("☁️ Salva nel cloud", width='stretch', disabled=not cloud_ready):
                    cards = [FlashCard(number=str(i+1),
                                      question=st.session_state.get(f"fc_ed_q_{i}", "").strip(),
                                      definition=st.session_state.get(f"fc_ed_d_{i}", "").strip())
                             for i in range(st.session_state.fc_editor_n)
                             if st.session_state.get(f"fc_ed_q_{i}", "").strip() or st.session_state.get(f"fc_ed_d_{i}", "").strip()]
                    if not cards:
                        st.warning("Aggiungi almeno una carta.")
                    else:
                        try:
                            file_stem = safe_name(fc_new_name)
                            json_path = cloud_path(cloud_settings["root"], fc_new_subject, "flashcard", f"{file_stem}.json")
                            cloud_upload_bytes(client, cloud_settings["bucket"], json_path, flashcards_to_json(cards), "application/json")
                            st.success(f"Mazzo '{fc_new_name}' salvato ({len(cards)} carte).")
                        except Exception as exc:
                            st.error(f"Errore: {exc}")
                if not cloud_ready:
                    st.caption("Configura il cloud per salvare.")

        with fc_sub_cloud:
            if not cloud_ready:
                st.info("Configura Supabase per accedere alla libreria cloud.")
            else:
                fc_cloud_subj = st.selectbox("Materia", options=list_subjects(client, cloud_settings) or ["(nessuna)"], key="fc_cloud_subject_sel")
                if fc_cloud_subj and fc_cloud_subj != "(nessuna)":
                    cloud_fc_files = list_flashcard_files(client, cloud_settings, fc_cloud_subj)
                    if not cloud_fc_files:
                        st.info("Nessuna flashcard per questa materia.")
                    else:
                        for cf_name in cloud_fc_files:
                            cf1, cf2, cf3 = st.columns([5, 1, 1])
                            with cf1:
                                st.markdown(f"{'🃏' if cf_name.endswith('.json') else '📄'} **{cf_name}**")
                            with cf2:
                                if st.button("Studia", key=f"fc_lib_open_{fc_cloud_subj}_{cf_name}", width='stretch'):
                                    try:
                                        for old in media_dir.glob("fc_*"): old.unlink(missing_ok=True)
                                        if cf_name.endswith(".json"):
                                            path = cloud_path(cloud_settings["root"], fc_cloud_subj, "flashcard", cf_name)
                                            cards = flashcards_from_json(cloud_download_bytes(client, cloud_settings["bucket"], path))
                                        else:
                                            cards, _ = load_flashcard_from_cloud(client, cloud_settings, fc_cloud_subj, cf_name, media_dir)
                                        st.session_state.fc_cards = cards
                                        st.session_state.fc_file_name = cf_name
                                        st.session_state.fc_subject = fc_cloud_subj
                                        st.session_state.fc_flipped = {}
                                        st.session_state.fc_prove_risposte = {}
                                        stato = load_flashcard_session(client, cloud_settings, fc_cloud_subj, safe_name(Path(cf_name).stem))
                                        st.session_state.fc_conosco = stato.get("conosco", [])
                                        st.session_state.fc_da_studiare = stato.get("da_studiare", [])
                                        st.rerun()
                                    except Exception as exc:
                                        st.error(f"Errore: {exc}")
                            with cf3:
                                if st.button("🗑️", key=f"fc_lib_del_{fc_cloud_subj}_{cf_name}", width='stretch'):
                                    try:
                                        if cf_name.endswith(".json"):
                                            path = cloud_path(cloud_settings["root"], fc_cloud_subj, "flashcard", cf_name)
                                            client.storage.from_(cloud_settings["bucket"]).remove([path])
                                        else:
                                            delete_flashcard_file(client, cloud_settings, fc_cloud_subj, cf_name)
                                        st.success("Eliminato.")
                                        st.rerun()
                                    except Exception as exc:
                                        st.error(exc)

        with fc_sub_file:
            st.markdown("Carica un DOCX o PDF con il formato flashcard.")
            with st.expander("📋 Formato atteso"):
                st.code("1. Termine\n[immagine opzionale]\n---\nDefinizione\n[immagine opzionale]\n\n2. Termine\n---\nDefinizione", language="text")
                st.markdown("Immagine **prima** del `---` → domanda. Immagine **dopo** il `---` → definizione.")
            fc_file_subj = st.text_input("Materia", value=st.session_state.selected_subject or "Generale", key="fc_file_subject")
            fc_upload = st.file_uploader("DOCX o PDF", type=["docx", "pdf"], key="fc_upload_file")
            bl, bc = st.columns(2)
            with bl:
                if st.button("📂 Apri localmente", disabled=fc_upload is None, width='stretch'):
                    try:
                        for old in media_dir.glob("fc_*"): old.unlink(missing_ok=True)
                        cards, _ = parse_flashcard_bytes(fc_upload.getvalue(), fc_upload.name, media_dir)
                        if not cards:
                            st.error("Nessuna flashcard trovata.")
                        else:
                            st.session_state.fc_cards = cards
                            st.session_state.fc_file_name = fc_upload.name
                            st.session_state.fc_subject = fc_file_subj
                            st.session_state.fc_flipped = {}
                            st.session_state.fc_prove_risposte = {}
                            if cloud_ready:
                                stato = load_flashcard_session(client, cloud_settings, fc_file_subj, safe_name(Path(fc_upload.name).stem))
                                st.session_state.fc_conosco = stato.get("conosco", [])
                                st.session_state.fc_da_studiare = stato.get("da_studiare", [])
                            else:
                                st.session_state.fc_conosco = []
                                st.session_state.fc_da_studiare = []
                            st.success(f"Caricate {len(cards)} flashcard.")
                            st.rerun()
                    except Exception as exc:
                        st.error(f"Errore: {exc}")
            with bc:
                if st.button("☁️ Carica nel cloud e apri", disabled=(fc_upload is None or not cloud_ready), width='stretch'):
                    try:
                        for old in media_dir.glob("fc_*"): old.unlink(missing_ok=True)
                        upload_flashcard_file(client, cloud_settings, fc_file_subj, fc_upload)
                        cards, _ = parse_flashcard_bytes(fc_upload.getvalue(), fc_upload.name, media_dir)
                        if cards:
                            st.session_state.fc_cards = cards
                            st.session_state.fc_file_name = fc_upload.name
                            st.session_state.fc_subject = fc_file_subj
                            st.session_state.fc_flipped = {}
                            st.session_state.fc_prove_risposte = {}
                            stato = load_flashcard_session(client, cloud_settings, fc_file_subj, safe_name(Path(fc_upload.name).stem))
                            st.session_state.fc_conosco = stato.get("conosco", [])
                            st.session_state.fc_da_studiare = stato.get("da_studiare", [])
                            st.success(f"Caricate {len(cards)} flashcard.")
                            st.rerun()
                    except Exception as exc:
                        st.error(f"Errore: {exc}")


# ============================================================
# DEBUG - Cloud e Session Status
# ============================================================

if st.checkbox("Mostra DEBUG - Cloud"):
    st.warning("MODALITA DEBUG ATTIVA")
    
    try:
        st.write("Cloud config:")
        st.write(f"- Bucket: {st.secrets['cloud']['bucket']}")
        st.write(f"- Root: {st.secrets['cloud']['root']}")
        st.write(f"- URL: {st.secrets['cloud']['supabase_url'][:40]}...")
        
        st.write("Cloud status:", "OK" if cloud_ready else "Non configurato")
        
        if cloud_ready:
            try:
                files = client.storage.from_(st.secrets['cloud']['bucket']).list(st.secrets['cloud']['root'])
                st.write(f"File nel cloud: {len(files) if files else 0}")
            except Exception as e:
                st.error(f"Errore nel leggere i file: {e}")
    except Exception as e:
        st.error(f"Errore nel caricamento config: {e}")

st.divider()

with st.expander("Diagnostica testo estratto"):
    if st.session_state.raw_text:
        st.text_area(
            "Testo riconosciuto",
            st.session_state.raw_text[:30000],
            height=300,
        )
    else:
        st.caption("Nessun testo estratto nella sessione corrente.")