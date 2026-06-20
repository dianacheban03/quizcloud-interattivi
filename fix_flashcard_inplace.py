"""
Modifica Flashcard.docx in-place preservando le immagini.
Operazioni:
  1. Intestazioni di sezione → svuotate (diventano blank tra card)
  2. Blank inseriti dove mancano tra card adiacenti
  3. Per ogni immagine che precede il suo '---', sposta il '---' PRIMA dell'immagine
     (le immagini devono stare sul lato risposta, dopo il separatore)
  4. Blank "interni" alla domanda → rimossi se senza immagini
  5. Separatori duplicati → rimossi
"""
import sys, io, re, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = r'C:\diana\POLI_new\AA_MAGISTRALE\Bionanotecnologie\Flashcard.docx'
DST = r'C:\diana\POLI_new\AA_MAGISTRALE\Bionanotecnologie\Flashcard_new.docx'

SEP_RE = re.compile(r'^\s*[-=—–]{1,}\s*$')
Q_RE   = re.compile(r'^\s*(\d+)\.\s+')

def is_section_header(t):
    if not t or SEP_RE.match(t) or Q_RE.match(t): return False
    letters = [c for c in t if c.isalpha()]
    if not letters or len(t) > 60: return False
    return sum(c.isupper() for c in letters) / len(letters) > 0.7

def has_image(p):
    return bool(p._p.findall('.//' + qn('a:blip')))

def is_blank(p):
    return not p.text.strip() and not has_image(p)

def make_blank_para():
    return OxmlElement('w:p')

def clear_para_text(p):
    """Svuota il testo del paragrafo senza toccare le immagini."""
    for r in p._p.findall(qn('w:r')):
        if not r.findall('.//' + qn('a:blip')):
            p._p.remove(r)

# ═══════════════════════════════════════════════════════════════════════
doc = Document(SRC)
paras = doc.paragraphs  # snapshot

# ── STEP 1+2: individua intestazioni e dove mancano blank tra card ──────
expanded = []
for i, p in enumerate(paras):
    for j, sub in enumerate(p.text.split('\n')):
        expanded.append((i, j, sub.strip()))

to_clear      = set()  # paragrafi da svuotare (erano intestazioni)
insert_before = set()  # inserire blank PRIMA di questo paragrafo

state = 'before'
last_main_q_num = 0
blank_before = False

for (para_idx, sub_idx, t) in expanded:
    if not t:
        # Solo un paragrafo vuoto reale (sub_idx==0) conta come blank_before;
        # le righe vuote da soft-return (sub_idx>0) NON costituiscono un vero blank.
        if sub_idx == 0:
            blank_before = True
        continue
    if is_section_header(t):
        to_clear.add(para_idx)
        blank_before = True
        last_main_q_num = 0
        continue
    if SEP_RE.match(t):
        if state == 'question':
            state = 'answer'
        blank_before = False
        continue
    m = Q_RE.match(t)
    if m:
        num = int(m.group(1))
        is_new = (num > last_main_q_num) and (
            blank_before or num == last_main_q_num + 1 or last_main_q_num == 0
        )
        if state == 'before':
            if is_new or last_main_q_num == 0:
                state = 'question'; last_main_q_num = num
        elif is_new:
            if not blank_before:
                insert_before.add(para_idx)
            state = 'question'; last_main_q_num = num
        blank_before = False
    else:
        blank_before = False

# ── STEP 3: sposta il '---' PRIMA delle immagini che lo precedono ────────
# Cerca: [IMG blank] → '---'   oppure   [IMG blank] → [blank*] → '---'
# In questi casi sposta il '---' PRIMA del primo [IMG blank]
# Esclude il caso in cui l'immagine è embedded nel testo (non blank para)
sep_to_move = {}   # {idx_del_sep: idx_da_inserire_prima}   (dove mettere il sep)

i = 0
while i < len(paras):
    p = paras[i]
    # Cerca un paragrafo immagine (testo vuoto + immagine)
    if has_image(p) and not p.text.strip():
        # Guarda all'indietro: c'è testo Q prima di questa immagine?
        # (non c'è già un separatore tra Q-text e questa immagine)
        has_sep_before = False
        for k in range(i - 1, -1, -1):
            pk = paras[k]
            if is_blank(pk): continue
            if SEP_RE.match(pk.text.strip()): has_sep_before = True
            break
        if not has_sep_before:
            # Guarda in avanti: il prossimo non-blank non-immagine è un '---'?
            for j in range(i + 1, min(i + 5, len(paras))):
                pj = paras[j]
                if has_image(pj): continue    # salta eventuali altri IMG
                if is_blank(pj): continue     # salta blank
                if SEP_RE.match(pj.text.strip()):
                    sep_to_move[j] = i   # sposta il sep da j a prima di i
                break
    i += 1

print(f'Separatori da spostare prima dell\'immagine: {sorted(sep_to_move.keys())} → target idx {[sep_to_move[k] for k in sorted(sep_to_move.keys())]}')

# ── STEP 4: individua blank interni ────────────────────────────────────
# Blank interno = blank che precede (direttamente o dopo immagini) un '---'
internal_blanks = set()
for i, p in enumerate(paras):
    if not is_blank(p):
        continue
    # Cerca il prossimo paragrafo non-blank, saltando immagini
    found_sep = False
    for j in range(i + 1, len(paras)):
        pj = paras[j]
        if is_blank(pj):
            continue           # altro blank: continua a cercare
        # Salta solo paragrafi SOLO-immagine (testo vuoto + immagine):
        # un paragrafo con testo E immagine è un Q o A reale, non saltarlo
        if has_image(pj) and not pj.text.strip():
            continue
        # primo para non-blank, non-solo-immagine
        if SEP_RE.match(pj.text.strip()):
            found_sep = True
        break
    if found_sep:
        internal_blanks.add(i)

print(f'Blank interni da rimuovere ({len(internal_blanks)}): {sorted(internal_blanks)}')
for i in sorted(internal_blanks):
    prev_t = paras[i-1].text.strip()[:50] if i > 0 else ''
    print(f'  [{i}] blank dopo: {repr(prev_t)}')

print(f'Intestazioni da svuotare ({len(to_clear)}): {sorted(to_clear)}')
print(f'Blank da inserire prima di: {sorted(insert_before)}')

# ── Applica modifiche ────────────────────────────────────────────────────
# Ordine inverso per non invalidare gli indici

# Prima: sposta i separatori (dal loro posto attuale a prima dell'immagine)
# Questo va fatto prima delle altre operazioni, lavorando su indici originali
# Raccogliamo gli elementi XML da spostare
sep_moves = []
for sep_idx, img_idx in sep_to_move.items():
    sep_para = paras[sep_idx]
    img_para = paras[img_idx]
    sep_moves.append((sep_para._p, img_para._p))

for sep_elem, img_elem in sep_moves:
    parent = sep_elem.getparent()
    if parent is None or img_elem.getparent() is None:
        continue
    # Rimuovi sep dalla posizione attuale
    parent.remove(sep_elem)
    # Inseriscilo prima di img_elem
    img_parent = img_elem.getparent()
    img_parent.insert(list(img_parent).index(img_elem), sep_elem)

# Ora le altre operazioni (su snapshot originale degli indici)
all_ops = []
for i in to_clear:
    all_ops.append((i, 'clear'))
for i in insert_before:
    all_ops.append((i, 'insert_blank'))
for i in internal_blanks:
    all_ops.append((i, 'delete'))
all_ops.sort(key=lambda x: x[0], reverse=True)

for (para_idx, action) in all_ops:
    p = paras[para_idx]
    parent = p._p.getparent()
    if parent is None:
        continue
    if action == 'clear':
        clear_para_text(p)
    elif action == 'insert_blank':
        parent.insert(list(parent).index(p._p), make_blank_para())
    elif action == 'delete':
        parent.remove(p._p)

# ── STEP 5: rimuovi separatori duplicati ────────────────────────────────
prev_was_sep = False
dups = []
for p in doc.paragraphs:
    t = p.text.strip()
    if not t: prev_was_sep = False; continue
    if SEP_RE.match(t):
        if prev_was_sep: dups.append(p._p)
        prev_was_sep = True
    else:
        prev_was_sep = False
print(f'Separatori duplicati rimossi: {len(dups)}')
for elem in dups:
    elem.getparent().remove(elem)

# ── Salva ────────────────────────────────────────────────────────────────
doc.save(DST)
print(f'\nSalvato: {DST}')

# ── Verifica con logica image-aware ─────────────────────────────────────
doc2 = Document(DST)
cur_q, cur_a, st, cards = [], [], 'q', []
for p in doc2.paragraphs:
    block = p.text.strip()
    img = bool(p._p.findall('.//' + qn('a:blip')))
    if not block and not img:
        if cur_q:
            cards.append((' '.join(cur_q)[:60], ' '.join(cur_a)[:40] if cur_a else '(vuota)'))
            cur_q, cur_a, st = [], [], 'q'
        continue
    for line in block.split('\n'):
        line = line.strip()
        if not line: continue
        if SEP_RE.match(line): st = 'a'; continue
        (cur_q if st == 'q' else cur_a).append(line)
if cur_q:
    cards.append((' '.join(cur_q)[:60], ' '.join(cur_a)[:40] if cur_a else '(vuota)'))

print(f'\nCard totali: {len(cards)}')
for i, (q, a) in enumerate(cards):
    print(f'  {i+1:2d}. {q}')
    print(f'      -> {a}')
